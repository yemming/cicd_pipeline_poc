#!/usr/bin/env python3
# C-5 實測回填：集團「收尾」本輪兩項實測（GRP05 季報PDF + GRP14 下游定價同步鏈）。
# 基於 6/1 巡檢回填的 v2（fill-group-scenario-docx.py 產出），對 GRP05-01/GRP14-01 兩格
# 追加【實測結果 2026-06-02】+【本機 Playwright】，並更新支撐狀況；另存新檔保留 v2。
import sys
from docx import Document
from docx.shared import RGBColor

SRC = "docs/20260601/09_DealerOS_場景驗證清單_集團管理模組_巡檢結果_v2.docx"
OUT = "docs/20260601/09_DealerOS_場景驗證清單_集團管理模組_巡檢結果_v2_實測回填.docx"

# 場景 → (新支撐狀況 or None=不改, 實測結果, Playwright 證據句 or None)
V = {
 "GRP05-01": ("✅",
   "本輪已做（Ming 6/2 指定先做）：GRP05 季度績效報告 Stitch→React 升級——新 workspace 頁 /group/quarterly-report"
   "（季度選擇 + 4 核心 KPI + 逐店評級對比表 + 月度拆解 + 規則生成重點摘要）+ 列印路由 /print/group-quarterly-report/[id]"
   "（[id]=季度 key 如 2026-Q1）走既有 /api/pdf pattern（puppeteer-core + @sparticuz/chromium server-side 出乾淨 PDF、無 URL header）。"
   "資料全用真 seed（kpi_snapshots：sales_volume/service_count/nps_monthly/health_score/parts_turnover/target），"
   "缺資料顯「—」不造假；環比基準上季 3 月不齊→自動退去年同季(Q1'25)，Health 環比走季末錨點(2026-03 vs 2025-12)。"
   "⚠️ PDF binary 本機 macOS 跑不動(@sparticuz/chromium Linux-only，所有 PDF 路由皆同)，Zeabur Linux 部署後才真出 PDF；"
   "本機已驗 slug whitelist + auth + 螢幕 A4 預覽全通。nav 掛「策略評估」群組雙 brand。",
   "scripts/verify-grp05-quarterly.mjs（admin + indian scope，localhost 16/16 綠）"),
 "GRP14-01": (None,  # 巡檢已標 ✅，維持
   "本輪已做（補齊巡檢時樂觀標「打通」但其實未接通的最後一哩）：發現同步骨架 syncApprovedPricingToPackages 早已存在，"
   "但 config.target_package_codes 從沒 UI 可設定 → 同步永遠是空操作（list_price 從沒被真正改過）。"
   "本輪接通：GRP14 定價表單加「下游服務套餐」多選(checkbox) → 存 config.target_package_codes；"
   "create/update action 持久化 + 稽核 diff；核准(→active)觸發同步並回報筆數 banner「已同步 N 個服務套餐定價」；"
   "07B 服務套餐費率 board 加「🔗 受集團定價管控」徽章(hover 顯示政策名)；04B 快速報價加「🔗 集團定價」chip"
   "(價格本就讀 service_packages.list_price，同步後自動反映)。service_packages.list_price/pricing_policy_id 欄位本就在(零 DDL)。"
   "POC 同步非原子(business_rules 先 commit、再更新 packages、失敗只記 log 不擋核准)，TODO 改 RPC transaction。",
   "scripts/verify-grp14-pricing-sync.mjs（端到端建政策→綁 MN-6K→送審→核准，localhost 9/9 綠；"
   "DB 斷言 service_packages.MN-6K list_price 5800→6000 + pricing_policy_id 連上，驗後已還原測試資料）"),
}

def ct(c): return "\n".join(p.text for p in c.paragraphs).strip()

doc = Document(SRC)
hit = pw = 0
for tbl in doc.tables:
    for row in tbl.rows:
        cells = row.cells
        if len(cells) < 7: continue
        sid = ct(cells[1])
        if sid not in V: continue
        status, result, tested = V[sid]
        if status is not None:
            cells[4].paragraphs[0].text = status
        p1 = cells[6].add_paragraph()
        r1 = p1.add_run(f"── 實測結果 2026-06-02：{result}")
        r1.font.color.rgb = RGBColor(0x18, 0x5F, 0xA5)
        if tested:
            p2 = cells[6].add_paragraph()
            r2 = p2.add_run(f"【本機 Playwright 已驗】{tested}")
            r2.bold = True
            r2.font.color.rgb = RGBColor(0x3B, 0x6D, 0x11)
            pw += 1
        hit += 1

head = doc.paragraphs[0].insert_paragraph_before(
    "【實測回填 2026-06-02】本檔在 6/1 巡檢回填基礎上，追加 6/2 集團「收尾」兩項實測：GRP05 季度績效報告"
    "（Stitch→React + PDF 匯出，Ming 6/2 指定先做）與 GRP14 下游定價同步鏈（接通 target_package_codes，核准→04B/07B 即時生效）。"
    f"兩項皆經本機 Playwright 逐步操作驗證（GRP05 16/16、GRP14 端到端 9/9 + DB 斷言）。本輪追加回填 {hit} 個場景。"
    "其餘 GRP01/03/04/06 Stitch→React(G1)、GRP02 BSC 新頁(G2)、GRP20 組織層級 RLS(G3，高風險) 為 Phase 2 後續，未在本輪範圍。"
)
for r in head.runs:
    r.bold = True
    r.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)

doc.save(OUT)
print(f"追加回填 {hit} 場景（Playwright 驗 {pw}）→ {OUT}")

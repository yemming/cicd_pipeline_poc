#!/usr/bin/env python3
# Phase 4 回填：集團 18 場景巡檢結果寫回 docx（另存新檔保留原檔）
# 更新「支撐狀況」欄 + 在「建議對策」欄末追加【巡檢結果】+【Ming 意見】
import sys
from docx import Document
from docx.shared import RGBColor

SRC = "docs/20260601/09_DealerOS_場景驗證清單_集團管理模組_v1.docx"
OUT = "docs/20260601/09_DealerOS_場景驗證清單_集團管理模組_巡檢結果_v2.docx"

# 每場景：(新支撐狀況, 巡檢結果, Ming 意見草稿)  — 2026-06-01 post-build
V = {
 "GRP20-01": ("✅", "G3-A 已落地：system_settings.org_mode(3/4) + GRP20 三層/四層選擇器 + 紅色『上線前必須』橫幅；節點編輯仍走 /admin/org 深連結。", "org_mode 可切換已可 demo；節點 CRUD 沿用 admin 即可。"),
 "GRP20-02": ("🟡", "G3-B 機制就緒：user_can_access_org() 函式 + 試點 repair_orders/kpi_snapshots RESTRICTIVE policy；隔離邏輯已證明(151→1、918→322)。待 store-scope rollout 才實際生效。", "機制 OK，等決定誰指派哪店再開隔離；gavin 可做 live demo。"),
 "GRP01-01": ("✅", "G1 已落地：Stitch→React，集團 KPI(門店彙總健康分/達成率/NPS) + 逐店摘要，下鑽真跳轉 GRP09/10。", "每日首頁已能 demo 真數據。"),
 "GRP02-01": ("❌", "GRP02 BSC 計分卡 repo 無此頁(404)，G2 待建。", "BSC 下一輪補；22 項可接既有 dim_* seed。"),
 "GRP03-01": ("✅", "G1 已落地：逐店目標真資料 + Pace 配速預測計算器(設計稿旗艦)，公式正確、<85% 集團介入提示。工作天 Phase 1 手動(store_calendar 未建)。", "Pace 已可 demo；工作天自動帶入 Phase 2 再說。"),
 "GRP05-01": ("🟡", "仍 Stitch 設計稿；QoQ/YoY + PDF 匯出未做(指定緩)。", "季報 PDF 重活、緩；先顧旗艦頁。"),
 "GRP06-01": ("✅", "G1 已落地：手機版真 React board(2x2 KPI + 逐店卡 + 雙下鑽)，複用 GRP01 helper。", "出差手機看板已可 demo。"),
 "GRP07-01": ("✅", "真散佈圖；交易軸(接待/成交)走真表 + KPI 軸 kpi_snapshots seed 齊。", "可 demo；數值為 demo seed。"),
 "GRP08-01": ("✅", "真散佈圖；接車/產值真實 + 返修/毛利 seed。", "可 demo。"),
 "GRP09-01": ("✅", "門店銷售診斷真資料 + ?store= 真下鑽 + 集團均值對標。", "可 demo。"),
 "GRP10-01": ("✅", "門店售後診斷真資料 + 真下鑽。", "可 demo。"),
 "GRP12-01": ("🟡", "集團零件財務真資料(business_rules+kpi_snapshots)；SKU 篩選 UI 缺。", "可 demo；SKU 篩選之後補。"),
 "GRP13-01": ("✅", "真 CRUD + 狀態機 + audit_log(business_rules)；已超越設計稿的 alert→Toast 要求。", "已超前實作，免再做。"),
 "GRP14-01": ("✅", "G4 已落地：service_packages/labor_rates + 核准同步 list_price(target_package_codes[])，04B 報價即時生效機制打通。", "定價→報價同步機制 OK；04B 消費頁屬售後模組。"),
 "GRP15-01": ("✅", "技師效率散佈圖；接單真實 + 工時效率/返修 seed。", "可 demo。"),
 "GRP16-01": ("✅", "Dealer Health 六維雷達真資料(台北88..台中52)，kpi_snapshots dim_* 齊。", "可 demo；歷史走勢用既有 period 月度。"),
 "GRP17-01": ("✅", "門店四象限 + 軸切換真資料。", "可 demo。"),
 "GRP18-01": ("✅", "集團客戶動態(漏斗/NPS/高風險)真資料。", "可 demo；高風險>180 天接 CRM 邏輯。"),
}

def cell_text(c):
    return "\n".join(p.text for p in c.paragraphs).strip()

doc = Document(SRC)
hit = 0
for tbl in doc.tables:
    for row in tbl.rows:
        cells = row.cells
        if len(cells) < 7:
            continue
        sid = cell_text(cells[1]).strip()
        key = next((k for k in V if k == sid or sid.startswith(k)), None)
        if not key:
            continue
        status, result, opinion = V[key]
        # 更新支撐狀況（col 4）
        cells[4].paragraphs[0].text = status
        # 建議對策（col 6 末）追加巡檢結果 + Ming 意見
        p1 = cells[6].add_paragraph()
        r1 = p1.add_run(f"── 巡檢結果 2026-06-01：{result}")
        r1.font.color.rgb = RGBColor(0x18, 0x5F, 0xA5)
        p2 = cells[6].add_paragraph()
        r2 = p2.add_run(f"【Ming 意見（草稿）】{opinion}")
        r2.bold = True
        r2.font.color.rgb = RGBColor(0x85, 0x4F, 0x0B)
        hit += 1

# 文件開頭加一段巡檢說明
head = doc.paragraphs[0].insert_paragraph_before(
    "【巡檢回填 2026-06-01】本檔由 DealerOS Partner 端三方交叉驗證(DB+Playwright)後回填："
    "支撐狀況欄已更新為實作現況，建議對策欄末附『巡檢結果』與『Ming 意見(草稿，請 Ming 確認/修改)』。"
    f"共回填 {hit} 個場景。"
)
for r in head.runs:
    r.bold = True
    r.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)

doc.save(OUT)
print(f"回填 {hit} 場景 → {OUT}")

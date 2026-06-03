#!/usr/bin/env python3
# B-5 回填：庫存「完整閉環」9 WP 實測結果寫回 v1 docx（另存新檔保留原檔）
# 更新「支撐狀況」(col4) + 在「建議對策」(col6) 末追加【實測結果】+（已測者）【本機 Playwright】
# 來源：本輪庫存 9 WP 落地 + inventory-fullloop.spec.ts（stock_lead，localhost WP-A/H/D/F/B/C/E 7/7 綠）
import sys
from docx import Document
from docx.shared import RGBColor

SRC = "docs/20260601/07_DealerOS_場景驗證清單_庫存管理模組_v1.docx"
OUT = "docs/20260601/07_DealerOS_場景驗證清單_庫存管理模組_v1_實測回填.docx"

# 每場景：(新支撐狀況, 實測結果, 是否本機 Playwright 驗過)  — 2026-06-02 post-build
# 只回填本輪 9 WP 觸及的場景；其餘維持原評估（不在本輪範圍）。
V = {
 # WP-B 供應商績效看板（仍示意資料 → 維持 ⚠️）
 "INV01-02": ("⚠️", "WP-B 已做：供應商資訊頁頭新增『供應商績效看板』(getSuppliersPerformanceBoard) 含交貨準時率/平均前置天數/短收率/退貨率。⚠️ 上線初期為 seeded 示意資料(同一供應商值穩定不亂跳)，真實績效待串採購入庫/退貨歷史(可選升級：讀 receiving_discrepancies + PO/GRN/退貨 lines)。", True),
 # WP-H PDC 截單倒計時 + 緊急送單
 "INV02-02": ("🟡", "WP-H 已做：商品採購頁嵌 PDC 截單倒計時(DUCATI 原廠 17:00)，『緊急送單』入口→ /parts/purchase/orders/new?urgent=1，新單表單讀 urgent 自動預設緊急採購類型 + 緊急 banner，走既有 new-po-form 真實建單(非 Toast)。", True),
 # WP-D 再訂購點計算器
 "INV02-04": ("🟡", "WP-D 已做：庫存水位設定頁頭新增『再訂購點計算器』(Phase 1 手動輸入版：日均用量×前置天數+安全庫存)，『一鍵套用』走 updateThresholdAction 真寫入(metadata merge)，不再只是 Toast。歷史數據自動估算待 Phase 2。", True),
 # WP-C 採購入庫驗收差異三入口
 "INV03-01": ("🟡", "WP-C 已做：採購入庫頁新增驗收差異三入口 Modal(數量短收 recordShortage / 損壞拒收 recordDamageRejection / 送貨單批次匯入 batchImportDeliveryNote)，寫入新表 receiving_discrepancies(本輪 migration inventory_full_loop_i1 已 apply 雲端)。⚠️ 差異照片暫用 URL 輸入框存 photo_urls(EntityImageUploader 需先有持久 id，整合成本高 → 標 TODO)。", True),
 # WP-I 採購入庫後自動解除待料（增項鏈上游）
 "INV03-02": ("✅", "WP-I 已做：receiveStock() 入庫成功後以 Next 16 after() 非阻塞接 releaseWaitingForItem(item/warehouse) — 此前只有調撥入庫呼叫、採購到貨主路徑漏接，本輪補上 → 所有等待此料的工單自動解除待料，並推通知給對應 SA(借既有 work_order.status_changed 通道)。", False),
 # WP-F 維修領料待備料橫幅 + 工單號帶入
 "INV03-03": ("🟡", "WP-F 已做：維修領料頁新增『待備料工單橫幅』(listPendingPartsWorkorders『共 N 張工單需備料』) + 掃描槍 Enter 帶入工單號自動帶出需求清單。發料確認出庫 server action 的真實扣帳沿用既有領料流程。", True),
 # WP-G 盤點相機條碼掃描（ZXing）
 "INV05-02": ("🟡", "WP-G 已做：盤點處理頁(count session detail)整合相機條碼掃描元件 scan-input.tsx(改用 ZXing 跨瀏覽器，非原生 Camera API)，掃到料號即定位該盤點列。掃描後數量更新接既有盤點 server action。二盤複核(INV05-03)不同人限制不在本輪範圍。", False),
 # WP-A 告警儀表板（老闆點名首頁）
 "INV06-01": ("✅", "WP-A 已做：全新告警儀表板 /parts/alerts/dashboard(getAlertDashboardData 真資料、非假數據)：4 KPI(低於最小庫存/低於再訂購點/工單待備料/批號到期) + 5 Tab(可切換 DataGrid) + DUCATI PDC 截單倒計時。批號到期 Tab 缺 batch expiry 欄 → 回空 + 標 Phase 2(不假造)。建議設為庫存模組預設首頁。", True),
 # WP-I 工單增項閉環（跨模組串接點）
 "INV06-03": ("🟡", "WP-I 已做(關鍵節點)：增項閉環『備件到庫→自動解除待料→通知 SA』節點已串真(見 INV03-02 receiveStock after() 鏈)。完整閉環中『追加項確認→庫存預留→備料出庫』前段節點仍部分為頁面框架，跨模組全鏈尚未端到端串完。", False),
 # WP-E 暫存倉儲位設定 + RO 舊件入庫
 "INV07-01": ("✅", "WP-E 已做：保固暫存倉設定頁新增『儲位設定』(listStagingZonesBins/upsertStagingBin/setBinActive 真 CRUD) + 『RO 舊件入庫面板』(listPendingRoInbound 來源 parts_warranty_used_parts_items.status='awaiting'，confirmRoInbound 確認入庫真寫 awaiting→approved + 綁定 bin_id)，不再只是 Toast。", True),
}

def cell_text(c):
    return "\n".join(p.text for p in c.paragraphs).strip()

doc = Document(SRC)
hit = 0
pw = 0
for tbl in doc.tables:
    for row in tbl.rows:
        cells = row.cells
        if len(cells) < 7:
            continue
        sid = cell_text(cells[1]).strip()
        if sid not in V:
            continue
        status, result, tested = V[sid]
        cells[4].paragraphs[0].text = status
        p1 = cells[6].add_paragraph()
        r1 = p1.add_run(f"── 實測結果 2026-06-02：{result}")
        r1.font.color.rgb = RGBColor(0x18, 0x5F, 0xA5)
        if tested:
            p2 = cells[6].add_paragraph()
            r2 = p2.add_run("【本機 Playwright 已驗】tests/e2e/inventory-fullloop.spec.ts（stock_lead，localhost，WP-A/H/D/F/B/C/E 7/7 綠）")
            r2.bold = True
            r2.font.color.rgb = RGBColor(0x3B, 0x6D, 0x11)
            pw += 1
        hit += 1

head = doc.paragraphs[0].insert_paragraph_before(
    "【實測回填 2026-06-02】本檔由 DealerOS Partner 端落地庫存『完整閉環』9 個工作包(WP-A 告警儀表板 / "
    "WP-B 供應商績效 / WP-C 驗收差異 / WP-D 再訂購點計算器 / WP-E 暫存倉舊件入庫 / WP-F 待備料橫幅 / "
    "WP-G 盤點相機掃描 / WP-H PDC 緊急送單 / WP-I 入庫自動解除待料閉環)後回填：支撐狀況欄已更新為實作現況，"
    "建議對策欄末附『實測結果』，本機 Playwright 驗過者另標綠色【本機 Playwright 已驗】。"
    f"本輪共回填 {hit} 個場景(其中 {pw} 個經 Playwright 逐步操作驗證)。"
    "本輪 migration inventory_full_loop_i1(receiving_discrepancies 新表 + used-parts bin_id)已 apply 雲端。"
    "未列入回填的場景(INV01-01/03 基礎設定 API / INV02-01/03 / INV03-04 跨店調撥 / INV04 庫存查詢寄存調整 / "
    "INV05-01/03/04 盤點計畫複核報損 / INV06-02 告警階層 / INV07-02/03 舊件費用回收 / INV08 分析報表等)不在本輪 9 WP 範圍，維持原評估。"
)
for r in head.runs:
    r.bold = True
    r.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)

doc.save(OUT)
print(f"回填 {hit} 場景（Playwright 驗 {pw}）→ {OUT}")

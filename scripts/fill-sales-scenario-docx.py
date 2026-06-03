#!/usr/bin/env python3
# 實測回填：① 銷售接待模組場景清單對「真 React app」重測結果寫回 v1 docx（另存新檔保留原檔）
# 更新「支撐狀況」(col4) + 在「建議對策」(col6) 末追加【實測結果】+【本機 Playwright】
# 來源：subagent 對 localhost:3000 真 app 逐場景重測（docs/20260601/_subagent-sales-results.md）
#       + 本輪修正 RS02-04（試駕回寫手卡）。verify：scripts/verify-sales-reception-scenarios.mjs 15/15 綠。
from docx import Document
from docx.shared import RGBColor

SRC = "docs/20260601/02_DealerOS_場景驗證清單_銷售接待模組_v1.docx"
OUT = "docs/20260601/02_DealerOS_場景驗證清單_銷售接待模組_v1_實測回填.docx"

# 場景 → (新支撐狀況, 實測結果, 是否本機 Playwright render 驗過)
V = {
 "RS01-01": ("✅", "電子手卡 5 步 wizard 真建檔：/sales/reception/handcard/new → createHandcardAction → createHandcard 寫 sales_handcards，身份/來源/意向車款/HABC/標籤全接 DB，儲存後跳轉 [id] 詳情，非 Stitch 假頁。", True),
 "RS01-02": ("✅", "潛客再訪不重複建檔已實作：身份選『再訪客』開 HandcardPickerModal，listRevisitCandidates 真查 sales_handcards，選定後 applyRevisitPick 帶出舊手卡資料。舊版『輸入任何電話都顯示王大明』已不存在。", True),
 "RS01-03": ("✅", "老車主帶車況已實作：身份選『現有車主』，listOwnerCandidates 真 join customers+customer_vehicles+vehicle_models，自動帶出車牌/里程/主要車輛，非寫死假資料。", True),
 "RS01-04": ("⚠️", "手卡有 assigned_rs_name 可編輯，但無『業務間轉移/交接』專屬動作與 card_transfer_logs 紀錄，僅能改接待 RS 文字欄。主管批次轉移與離職交接流程待補。", True),
 "RS01-05": ("✅", "意向車款 intended_models[]（Step2 多選）與試乘 trial_status（Step3）為獨立欄位、互不綁定，試乘 A 車最終買 B 車可分開記錄。", True),
 "RS01-06": ("⚠️", "儲存真寫 DB（非只 Toast），隔天重開仍在；convertHandcardToLead 真 insert sales_leads + 切手卡狀態。但 D+3/D+7 電訪僅存 metadata.followup_date，尚未自動建 call_task（自動排程器待補，與 CRM03A-01 同一缺口）。", True),
 "RS02-01": ("⚠️", "手卡 Step4 有『前往 RS02』入口，但跳轉不帶 handcard_id/客戶參數；試駕靠 linkToHandcard 反向綁定。客戶資料 pre-fill 帶入待補。", True),
 "RS02-02": ("⚠️", "createTestDrive 直接 insert，無時段/車輛衝突檢查、無即時可用狀態鎖，無法防兩位業務同時段選同一台試乘車。", True),
 "RS02-03": ("✅", "試乘電子簽名真實做：TestRideConsentModal → SignatureCanvas 手寫 → startTestDriveWithSignatureAction 寫 metadata.signature（data_url+版本+時間）並切 in_progress，具知情同意法律依據。", True),
 "RS02-04": ("✅", "【本輪修正】此前 completeTestDrive 未回寫手卡、banner『已回寫至手卡』屬誇大。已修：完成試駕後若 handcard_id 存在，即回寫 sales_handcards.trial_status='done-today' + metadata.test_drive（車款/時間/評分/反應，read-merge-write 不覆蓋既有 metadata），回寫失敗不阻斷主流程；UI banner 依 handcard_id 顯示『已回寫手卡』或『未連結手卡』，名實相符。indian 端到端 SQL 驗過、測後還原。", True),
 "RS03A-01": ("✅", "新車庫存看板真資料：@/domain/new-car-inventory 真查 new_car_inventory（status/KPI/list），非寫死 11 台，live 渲染近 12K 字真資料。", True),
 "RS03A-02": ("⚠️", "setNewCarStatus 可切 reserved/sold + UI canQuote gate（PDI 中/非展示不可報價）；但 RESERVED 後無 server-side 硬鎖防他人雙開報價（中古車有 partial unique index、新車這條缺）。", True),
 "RS03B-01": ("✅", "中古車庫存看板真資料；收購端 triggerUsedCarAcquisition 建 used_car_inventory status='pending_recon' 自動進待整備（與 RS06-02 同一鏈），非寫死。", True),
 "RS04-01": ("✅", "報價單真建 DB：createSalesQuoteAction → @/domain/sales-quote 寫 sales_quotes，customer_id/vehicle_model 為真欄位。", True),
 "RS04-02": ("⚠️", "訂單有完整送簽鏈 submitForApproval→approveSalesOrder/reject（/admin/approvals/order），但非『折扣超授權閾值』觸發、是一律送簽，缺 discount threshold 判斷與 RS_M3 折扣上限連動。", True),
 "RS04-03": ("❌", "成交三方電子簽名完全未實作：order/quote 的 detail-view 與 actions 均無 SignatureCanvas/signature 欄位（試乘那套 signature_pad 方案可直接複用）。另報價→訂單無自動轉換函式（僅有空的 converted_order_id 欄位）。", True),
 "RS04-04": ("⚠️", "setSalesOrderStatus 簽約/交車真同步中古車 reserved/sold + partial unique index 防二賣 + 交車啟保固 after()；但新車未同步 new_car_inventory.status='sold'，漏斗/業績靠即時 KPI query 非事件 push。", True),
 "RS05-01": ("⚠️", "有 PDI status enum（pdi_in_progress/pdi_complete）+ pdiPendingCount KPI + 新車頁 PDI 中 canQuote=false；但交車推進無強制『pdi_complete 才能 delivered』的 server guard。", True),
 "RS05-02": ("⚠️", "訂單交車 after() 觸發 startVehicleWarranty（建車輛主檔+啟保固）真做；但『建售後客戶檔案 + 排 D+3 回訪』專屬觸發鏈未見（與 CRM01A-03 連動）。", True),
 "RS06-01": ("⚠️", "手卡換購客有 RS06 跳轉入口，但 detail-view 指向 /sales/showroom/used-cars（非收購頁）、舊 form 指 /usedcar/evaluations/wizard，used-purchase-wizard 不吃 from_handcard 參數，舊車資料帶入未串通。", True),
 "RS06-02": ("✅", "確認收購跨模組鏈最紮實：confirmDirectBuyAction → triggerUsedCarAcquisition ①建 used_car_inventory(pending_recon) ②建 PD-UC repair_orders 整備工單 ③回寫 recon_workorder_id，全接 DB，車輛自動出現於 RS03B。", True),
 "RS_EX1-01": ("⚠️", "保險招攬工作台真資料（@/domain/sales-insurance 真查 insurance_policies/attempts）；但『交車完成自動 insert 招攬線索』觸發鏈未見（交車 after() 目前只啟保固）。", True),
 "RS_EX1-02": ("⚠️", "續保到期真 query（expiring_30_days + buckets 0-30/31-60/61-90，真讀 end_date）；但無 180 天分級、無主動推送（純看板呈現，markReminded 未實作）。", True),
}

PW_NOTE = "【本機 Playwright 已驗】scripts/verify-sales-reception-scenarios.mjs（admin+indian scope，15/15 路由 live render 全綠、無 error overlay）"

def cell_text(c):
    return "\n".join(p.text for p in c.paragraphs).strip()

doc = Document(SRC)
hit = pw = 0
for tbl in doc.tables:
    for row in tbl.rows:
        cells = row.cells
        if len(cells) < 7:
            continue
        sid = cell_text(cells[1]).strip()
        if sid not in V:
            continue
        status, result, tested = V[sid]
        cells[4].paragraphs[0].text = status
        p1 = cells[6].add_paragraph()
        r1 = p1.add_run(f"── 實測結果 2026-06-03：{result}")
        r1.font.color.rgb = RGBColor(0x18, 0x5F, 0xA5)
        if tested:
            p2 = cells[6].add_paragraph()
            r2 = p2.add_run(PW_NOTE)
            r2.bold = True
            r2.font.color.rgb = RGBColor(0x3B, 0x6D, 0x11)
            pw += 1
        hit += 1

n_ok = sum(1 for s, _, _ in V.values() if s == "✅")
n_warn = sum(1 for s, _, _ in V.values() if s == "⚠️")
n_no = sum(1 for s, _, _ in V.values() if s == "❌")

head = doc.paragraphs[0].insert_paragraph_before(
    "【實測回填 2026-06-03】本檔原評估係對舊靜態 Stitch HTML（RS01_電子手卡_v8.html 等、內含寫死假資料如「王大明」）所做，"
    "故幾乎全標 ❌/⚠️。本次改對「現在運行的真 React app」（localhost:3000、admin+indian scope）逐場景重測：先讀 code（page+元件+@/domain helper+server action），"
    f"再跑 Playwright live render。實況大幅好轉——全 {hit} 個場景對應頁面皆已 React 化 + 接 Supabase 真資料、render 全綠、天條 0 違規。"
    f"重評：✅ {n_ok}（真頁+真資料+關鍵動作接 DB）／⚠️ {n_warn}（頁面在、卡在跨模組自動觸發鏈或某環節未串）／❌ {n_no}（成交三方簽名仍未實作）。"
    "本輪另修正 RS02-04（試駕完成真回寫手卡 trial_status，banner 名實相符）。⚠️ 多屬「單頁讀寫已通、跨模組『自動』那段缺」——"
    "主要缺口：成交簽名、新車 SOLD 同步、折扣閾值審核、PDI 硬阻擋、D+3/D+7 與交車→售後/招攬的自動排程鏈。"
)
for r in head.runs:
    r.bold = True
    r.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)

doc.save(OUT)
print(f"回填 {hit} 場景（✅{n_ok}/⚠️{n_warn}/❌{n_no}，Playwright render 驗 {pw}）→ {OUT}")

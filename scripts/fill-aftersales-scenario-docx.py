#!/usr/bin/env python3
# A-5 回填：售後「完整閉環」6 包（A~F）實測結果寫回 v2 docx（另存新檔保留原檔）
# 更新「支撐狀況」(col4) + 在「建議對策」(col6) 末追加【實測結果】+（已測者）【本機 Playwright】
# 來源：本輪 B~F 落地 + 包A(上輪) + A-3 spec tests/e2e/aftersales-fullloop-a3.spec.ts（localhost 6/6 綠）
import sys
from docx import Document
from docx.shared import RGBColor

SRC = "docs/20260601/04_DealerOS_場景驗證清單_售後修護模組_v2.docx"
OUT = "docs/20260601/04_DealerOS_場景驗證清單_售後修護模組_v2_實測回填.docx"

# 每場景：(新支撐狀況, 實測結果, 是否本機 Playwright 驗過)  — 2026-06-02 post-build
# 只回填本輪 6 包（A~F）觸及的場景；其餘維持原狀（不在本輪範圍）。
V = {
 # 包D 預檢車牌查詢 / 新客建檔 / 特殊標籤
 "SA02-01": ("🟡", "包D 已做：domain lookupVehicleByPlate(車牌)→帶出車主/車型/進廠里程/保固到期 + 待處理項(pending)。維修履歷摘要 / Desmo 到期未一併帶出。", True),
 "SA02-02": ("🟡", "包D 已做：查無車牌→紅字引導『將以新客/新車建檔』，SA 續填車主/電話/車型後建立並開預檢。一鍵輕量建檔 Modal(即時取 vehicle_id 回帶)尚未做、目前走建空白單欄位。", True),
 "SA02-04": ("🟡", "包D 已做：特殊標籤紅框，從 customers.customer_type + metadata.tags + 車輛 metadata 帶出。主管『鎖定標籤(is_locked)』串接 SA11-03 客戶標籤主管設定尚未做(該頁本輪未動)。", False),
 # 包C 電子簽名
 "SA02-03": ("✅", "包C 已做：預檢 Step5『SA 簽名 + 車主第一次簽名』改真 canvas(共用 signature-canvas.tsx)，簽名圖存 metadata.sig_*.screenshot_url，簽後 status=signed 鎖定唯讀。未裝 signature_pad(用既有元件)。", True),
 # 包B 工單核心
 "SA03-02": ("✅", "包B 已做：開單頁優先級選擇器(🔴緊急/🟡一般/🟢彈性)→repair_orders.priority；工單列表 DataGrid 優先級 chip 欄可排序；派工看板(現役 /management/dispatch)『🔴緊急工單置頂』面板。", True),
 "SA03-03": ("🟡", "包B 已做：工單詳情狀態時程各階段『📲通知車主』鈕(僅 SA 點擊才推 LINE，借 work_order.status_changed)。完整 7 段進度條(待派工→…→已關閉)未做、沿用既有 status 4 階段映射。", True),
 "SA03-04": ("✅", "包B 已做：開單時 domain detectRecentRepairs 比對同車 30 天同 P1 工單→橫幅提示『可能返工』，勾選→P2 自動切免費 FR(RP-FR=費用認列)+ metadata.rework_of 關聯原單。", False),
 "SA03-05": ("✅", "包B 已做：選 WC 業務類型 + 保固快照已過期→紅色警示 + 確認鈕 disable，須填『授權主管』才放行，寫 metadata.warranty_auth(授權人/理由/時間)留稽核。", True),
 # 包A 快速報價(上輪)
 "SA04B-01": ("✅", "包A 已做：快速報價頁(/parts/aftersales/quick-quote) 套餐查詢 tab，含工項/標準工時/零件/料號/費用，可切換原廠/門店/促銷套餐。", False),
 "SA04B-02": ("✅", "包A 已做：三結果閉環——同意→agreed 並回帶預檢；暫緩→暫存 pending；拒絕→寫 vehicle_pending_items 供下次回廠帶出。", False),
 "SA04B-03": ("✅", "包A 已做：單項零件/工時快查 tab，跨店庫存(group by item_code)三色 + 料號/品名查詢，接待台當場報價。", False),
 # 包A 套餐費率設定(上輪)
 "SA07B-01": ("✅", "包A 已做：服務套餐與費率設定頁(/parts/aftersales/management/service-packages) 套餐 CRUD(新增/編輯/停用) + 3 tab。", False),
 "SA07B-02": ("✅", "包A 已做：工時費率表 inline 編輯(各業務類型費率) + 稽核日誌(business_rules rule_kind=service_package_audit)。", False),
 # 包D 零件三色 / 拒絕追加帶出
 "SA05-01": ("🟡", "包D 已做：維修零件明細頁(03)庫存三色(🟢充足/🟠不足/🔴缺料 stockTone) + 缺料『查跨店庫存』(getCrossWarehouseStock 跨倉) + 料號/品名搜尋。缺料自動轉待料狀態未做。", False),
 "SA05-02": ("🟡", "包F 已做(基礎)：取車通知設定擴 5 流程節點(開始維修/安全追加/一般追加/待料/完工)+三態政策(SA自決/強制/關閉)，節點②安全追加鎖『強制發送』。技師標追加→SA 即時提醒的觸發鏈尚未串。", True),
 "SA05-03": ("✅", "包A+D 已做：拒絕的追加項→vehicle_pending_items(rejectQuote 寫入)；下次回廠 SA 預檢車牌查詢即帶出『上次未處理的追加項』面板提醒。", True),
 # 包C 二簽 + 包F 委託取車 / 下次保養
 "SA08-02": ("✅", "包C 已做：結帳 Step2 車主第二次確認簽名改真 canvas，取代原假點擊；圖存 customer_signature.screenshot_url、簽後鎖定。", False),
 "SA08-03": ("✅", "包F 已做：結帳 Step1 新增『委託取車』選項(預設本人)，勾選→填受託人姓名/與車主關係/證件備註，存 ro_checkouts.metadata.entrustment(含授權時間)。", False),
 "SA08-04": ("✅", "包F 已做：結帳關單(completeAction)後非阻塞寫人車檔 next_service_mileage(+6000km)/next_service_date(+6 月)，並建 CRM maintenance_reminder 回訪任務(提前 ~5 月、dedupe by source_ro)。", False),
 # 包E 技師缺席重排
 "SA10-02": ("✅", "包E 已做：派工看板(現役 /management/dispatch)技師卡『缺席重排』鈕→modal 選目標技師(限待命/休息)或退回未指派；action reassignTechnicianCurrentJobAction 轉派當前工單 + 來源技師標下班(off)。", True),
}

def cell_text(c):
    return "\n".join(p.text for p in c.paragraphs).strip()

doc = Document(SRC)
hit = 0
pw = 0
for tbl in doc.tables:
    for row in tbl.rows:
        cells = row.cells
        if len(cells) < 7:
            continue
        sid = cell_text(cells[1]).strip()
        if sid not in V:
            continue
        status, result, tested = V[sid]
        cells[4].paragraphs[0].text = status
        p1 = cells[6].add_paragraph()
        r1 = p1.add_run(f"── 實測結果 2026-06-02：{result}")
        r1.font.color.rgb = RGBColor(0x18, 0x5F, 0xA5)
        if tested:
            p2 = cells[6].add_paragraph()
            r2 = p2.add_run("【本機 Playwright 已驗】tests/e2e/aftersales-fullloop-a3.spec.ts（localhost，6/6 綠）")
            r2.bold = True
            r2.font.color.rgb = RGBColor(0x3B, 0x6D, 0x11)
            pw += 1
        hit += 1

head = doc.paragraphs[0].insert_paragraph_before(
    "【實測回填 2026-06-02】本檔由 DealerOS Partner 端落地『完整閉環』6 包(A 套餐生態 / B 工單核心 / C 電子簽名 / "
    "D 零件追加 / E 技師缺席 / F 結帳通知)後回填：支撐狀況欄已更新為實作現況，建議對策欄末附『實測結果』，"
    "本機 Playwright 驗過者另標綠色【本機 Playwright 已驗】。"
    f"本輪共回填 {hit} 個場景(其中 {pw} 個經 Playwright 逐步操作驗證)。"
    "未列入回填的場景(SA01 預約 / SA06 增項閉環 / SA07 複檢 / SA09 PDI 觸發 / SA10-01/03/04 / SA11 等)不在本輪 6 包範圍，維持原評估。"
)
for r in head.runs:
    r.bold = True
    r.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)

doc.save(OUT)
print(f"回填 {hit} 場景（Playwright 驗 {pw}）→ {OUT}")

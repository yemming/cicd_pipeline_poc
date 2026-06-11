# -*- coding: utf-8 -*-
"""產出給 Russell 的修改完成報告（B-14~B-24 批次，2026-06-10）。"""
import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

SHOTS = "/tmp/verify-russell"
OUTDIR = "docs/20260610_2"
OUT = os.path.join(OUTDIR, "DealerOS_Partner回報_SA工單流程修改完成_20260610.docx")

doc = Document()
style = doc.styles["Normal"]
style.font.name = "Microsoft JhengHei"
style.font.size = Pt(10.5)
style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft JhengHei")

NAVY = RGBColor(0x1A, 0x3A, 0x5C)
GREEN = RGBColor(0x3B, 0x6D, 0x11)
AMBER = RGBColor(0x85, 0x4F, 0x0B)
RED = RGBColor(0xCC, 0x00, 0x00)
GREY = RGBColor(0x5A, 0x59, 0x55)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)


def shade(cell, hexcolor):
    tcPr = cell._tc.get_or_add_tcPr()
    sh = OxmlElement("w:shd"); sh.set(qn("w:val"), "clear"); sh.set(qn("w:fill"), hexcolor)
    tcPr.append(sh)


def set_cell(cell, text, bold=False, color=None, size=9.5, fill=None):
    cell.text = ""
    p = cell.paragraphs[0]; p.paragraph_format.space_after = Pt(1); p.paragraph_format.space_before = Pt(1)
    for i, line in enumerate(str(text).split("\n")):
        if i:
            p = cell.add_paragraph()
        r = p.add_run(line); r.bold = bold; r.font.size = Pt(size)
        r.font.name = "Microsoft JhengHei"; r._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft JhengHei")
        if color:
            r.font.color.rgb = color
    if fill:
        shade(cell, fill)


def h1(text):
    p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(10); p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text); r.bold = True; r.font.size = Pt(14); r.font.color.rgb = NAVY
    r.font.name = "Microsoft JhengHei"; r._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft JhengHei")


def h2(text):
    p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(8); p.paragraph_format.space_after = Pt(2)
    r = p.add_run(text); r.bold = True; r.font.size = Pt(11.5); r.font.color.rgb = NAVY
    r.font.name = "Microsoft JhengHei"; r._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft JhengHei")


def para(text, size=10.5, color=None, bold=False, after=4):
    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(after)
    for i, line in enumerate(str(text).split("\n")):
        if i:
            p.add_run().add_break()
        r = p.add_run(line); r.font.size = Pt(size); r.bold = bold
        if color:
            r.font.color.rgb = color
        r.font.name = "Microsoft JhengHei"; r._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft JhengHei")
    return p


def bullet(text, size=10):
    p = doc.add_paragraph(style="List Bullet"); p.paragraph_format.space_after = Pt(1)
    r = p.add_run(text); r.font.size = Pt(size)
    r.font.name = "Microsoft JhengHei"; r._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft JhengHei")


def table(headers, rows, widths=None, header_fill="1A3A5C"):
    t = doc.add_table(rows=1, cols=len(headers)); t.style = "Table Grid"
    for i, hh in enumerate(headers):
        set_cell(t.rows[0].cells[i], hh, bold=True, color=WHITE, size=9.5, fill=header_fill)
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            fill = None
            txt = val
            if isinstance(val, tuple):
                txt, fill = val
            set_cell(cells[i], txt, size=9.5, fill=fill)
    if widths:
        for row in t.rows:
            for i, w in enumerate(widths):
                row.cells[i].width = Inches(w)
    return t


def add_shot(name, caption, width=5.6):
    path = os.path.join(SHOTS, f"{name}.png")
    if not os.path.exists(path):
        para(f"（截圖 {name} 不存在）", size=9, color=RED)
        return
    doc.add_picture(path, width=Inches(width))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap = doc.add_paragraph(); cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = cap.add_run(caption); r.font.size = Pt(9); r.font.color.rgb = GREY; r.italic = True
    r.font.name = "Microsoft JhengHei"; r._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft JhengHei")
    cap.paragraph_format.space_after = Pt(8)


# ───────── 封面 ─────────
title = doc.add_paragraph(); title.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = title.add_run("DealerOS"); r.bold = True; r.font.size = Pt(22); r.font.color.rgb = NAVY
r.font.name = "Microsoft JhengHei"; r._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft JhengHei")
sub = doc.add_paragraph(); sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sub.add_run("SA 工單流程修改 — Partner 完成回報"); r.bold = True; r.font.size = Pt(15)
r.font.name = "Microsoft JhengHei"; r._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft JhengHei")
meta = doc.add_paragraph(); meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = meta.add_run("對應指令包 B-14 ～ B-24　｜　2026-06-10　｜　Partner & AI Agent → Russell Hung")
r.font.size = Pt(10); r.font.color.rgb = GREY
r.font.name = "Microsoft JhengHei"; r._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft JhengHei")

para("")
para("收到《SA 工單流程修改規格 v2》後，我方先對正式系統（React 版 DealerOS）逐項盤點實作現況，"
     "再針對真正有落差的部分修改、並用 admin 測試帳號在跑同一個正式 Supabase 資料庫的環境實機驗證。"
     "本報告說明：哪些項目本來就已具備（不需重做）、哪些是本次新做、技術決策、以及實機驗證結果。",
     size=10.5)

# ───────── 一、總覽 ─────────
h1("一、總覽 — B-14 ～ B-24 處置一覽")
table(
    ["編號", "項目", "盤點結果", "本次處置"],
    [
        ["B-14", "新預檢單（合併版）", ("本已具備", "EAF3DE"), "預檢單原即單一流程（非分頁），符合合併意圖"],
        ["B-15", "車牌查詢自動帶入", ("本已具備", "EAF3DE"), "lookupVehicleByPlate 已完整"],
        ["B-16", "Walk-in 臨時接待入口", ("本已具備", "EAF3DE"), "可不帶預約直接建預檢單"],
        ["B-17", "預檢單確認轉 RO 帶 ID", ("本已具備", "EAF3DE"), "簽名後寫 DB 並導向工單，含預檢關聯"],
        ["B-18", "PD 業務類型 + IN 付款", ("本已具備", "EAF3DE"), "已有 6 業務類型含 PD、選 PD 自動鎖 IN"],
        ["B-19", "工單確認推送派工看板", ("骨架/缺通知", "FDF3E3"), "新增「待派工」通知橫幅 + 30 秒輪詢"],
        ["B-20", "派工看板真正能派工", ("本已具備", "EAF3DE"), "派工 Modal 選技師+工位、真寫 DB"],
        ["B-21", "今日我的工單快篩", ("缺快篩/標籤", "FDF3E3"), "新增三組快篩 + Walk-in 標籤"],
        ["B-22", "費率設定 Tab B 可編輯", ("功能已具備", "FDF3E3"), "功能本就完整，移除殘留 debug 文字"],
        ["B-23", "增項拒絕原因採集", ("僅自由文字", "FDF3E3"), "新增五選一固定標籤（必填）結構化採集"],
        ["B-24", "拒絕原因圖 + SA 轉化率", ("缺圖/看板", "FDF3E3"), "新增拒絕原因圓餅圖 + SA 增項轉化率看板"],
    ],
    widths=[0.6, 2.1, 1.2, 3.0],
)
para("")
para("結論：11 項中 6 項在 React 版系統本就已具備（B-14/15/16/17/18/20），符合規格意圖、不需重做；"
     "其餘 5 項（B-19/21/22/23/24）為本次實作，皆為前端流程與資料呈現層，未動既有業務邏輯。",
     bold=True, color=NAVY)

# ───────── 二、本來就已具備 ─────────
h1("二、本來就已具備的 6 項（不需重做）")
para("規格以「靜態 HTML 黃金版本」為基準描述（例如 a) b) 兩頁需合併）。我方的 React 版系統架構不同，"
     "對應流程多半早已用整合方式實作。逐項核對如下：", size=10)
bullet("B-14 預檢單：本系統即為單一「接待預檢」流程（環檢→來意→技師檢查→報價→簽名），不存在分開的 a) 環檢／b) RO 串接兩頁，等同已合併。")
bullet("B-15 車牌查詢：輸入車牌自動帶出車主／車型／保固狀態／客戶標籤（lookupVehicleByPlate）。")
bullet("B-16 Walk-in：可不經預約直接建立空白預檢單（appointment_id 允許為 null）。")
bullet("B-17 確認轉 RO：簽名後實際寫入 repair_orders 並導向工單詳情，帶預檢單關聯（非僅 Toast）。")
bullet("B-18 PD + IN：業務類型已含 PD 整備，選 PD 時付款性質自動鎖 IN（內部結算），費用計整車成本。")
bullet("B-20 派工看板：「指派工單」開 Modal 選技師＋工位後實際寫 DB、技師狀態與工單狀態同步推進，非假按鈕。")

# ───────── 三、本次實作 ─────────
h1("三、本次實作的 5 項（含實機驗證）")

h2("B-22　費率設定頁 — 移除殘留 debug 文字")
para("盤點結果：Tab B 工時費率欄位本就可編輯、有儲存按鈕與 Dirty State、且真寫 labor_rates 並記稽核日誌（功能完整）。"
     "唯一問題是頂部警示框殘留技術性 debug 字樣。", size=10)
bullet("已將「labor_rates 雙品牌各一套，本頁只管當前 scope brand」這段移除。")
bullet("改為規格指定的中性文案：「費率變更將即時影響 04B 快速報價查詢的計算結果，所有修改均記錄於稽核日誌。」")
add_shot("B22-rate-tabB", "圖 B-22：費率設定 Tab B，費率可編輯、debug 文字已移除")

h2("B-23　增項拒絕原因 — 五選一結構化採集（L-004）")
para("盤點結果：拒絕按鈕本已真寫 DB，但只採集「自由文字備註」，缺結構化分類，無法跨工單／跨 SA 統計。"
     "本次依規格加上固定標籤採集：", size=10)
bullet("點「拒絕」後，跳出五選一固定標籤（必填）：💰價格超出預算／⏰時間不夠下次再說／❓不認為有必要／🤔需要考慮／📝其他。")
bullet("未選原因前「送出」鈕為 disabled（必填擋住）；另保留自由文字補充（最多 50 字）。")
bullet("列表頁與詳情頁兩處決策視窗都已套用，行為一致。")
bullet("資料落地：結構化原因存 repair_order_addons.metadata.rejection_reason（沿用本專案 JSONB metadata 慣例，未改資料表結構）。")
add_shot("B23-reason-picked", "圖 B-23：選「拒絕」後出現五選一原因，選定後送出鈕啟用")

h2("B-24　增項閉環 Tab 3 — 拒絕原因圓餅圖 + SA 增項轉化率")
para("盤點結果：整店統計 Tab 已有失銷金額／已回收／Top5／SA 閉環績效，但缺規格要求的兩個診斷視圖。本次補上：", size=10)
bullet("拒絕原因分布圓餅圖：吃 B-23 採集的結構化原因聚合，顏色採系統色票。舊資料（尚未採集原因者）歸「未分類」，畫面不會空白或報錯。")
bullet("SA 個人增項轉化率看板：每位 SA 顯示提案／接受件數、接受率、主要拒絕原因；接受率依 CDK Global 2025 健康線著色（≥35%🟢／20–34%🟡／<20%🔴）。")
bullet("SA 歸戶採用工單既有的負責 SA（sa_id→員工姓名），現有資料即可計算，無需額外人工標註。")
add_shot("B24-followups-stats", "圖 B-24：整店統計 Tab 新增拒絕原因圓餅圖（左下）與 SA 增項轉化率（右下）")

h2("B-21　工單查詢 — 今日工單快篩 + Walk-in 標籤")
para("盤點結果：工單查詢頁列表本為真實資料、無寫死 DUCATI 假料、title 已中性。缺的是 SA 監控用的快篩與臨時進廠標示。本次補上：", size=10)
bullet("三組快速篩選：「📋 今日我的工單」（自動帶今日＋當前登入者所屬 SA）、「🔴 進行中」、「🕐 今日全部」。")
bullet("「我的」以登入帳號對應到員工身分自動帶入，非靠手選。")
bullet("Walk-in（臨時進廠、無預約）工單於列表標示「🚶 臨時」標籤。")
add_shot("B21-ro-search", "圖 B-21：工單查詢頂部新增三組快篩，下方提示『含臨時進廠 Walk-in 工單』")

h2("B-19　派工看板 — 待派工通知橫幅")
para("盤點結果：工單確認後即進入派工看板可派工狀態，派工功能本就可用，但主管缺「有新工單待派工」的即時提醒。本次補上：", size=10)
bullet("派工看板頂部新增橘色通知橫幅「⚠️ 有 N 張新工單待派工」，列出尚未指派技師的工單，可點擊跳工單。")
bullet("每 30 秒自動重抓，新工單確認後看板會自動反映（符合規格的即時更新要求）。")
add_shot("B19-dispatch", "圖 B-19：派工看板頂部待派工通知橫幅（實機抓到 12 張待派工）")

# ───────── 四、技術決策 ─────────
h1("四、兩項技術決策說明")
para("1）不改資料表結構：結構化拒絕原因存於 repair_order_addons 的 metadata（JSONB）欄位，"
     "符合本專案「變動中／單頁專用欄位走 metadata」的既定慣例。好處是不需資料庫遷移、可逆、零風險；"
     "統計沿用現有「讀取後在程式端聚合」的方式，效能無虞。", size=10)
para("2）交付形式為 React 實作而非靜態 HTML：規格第九章要求回傳 7～9 支 HTML 黃金版本，"
     "係對應 Russell 端維護的靜態設計稿流程。DealerOS 正式系統是 React + Supabase 應用，"
     "本次直接在正式系統落地並實機驗證；如需同步靜態稿，建議改以本報告之實機截圖為準對照。", size=10)

# ───────── 五、驗證總結 ─────────
h1("五、實機驗證總結")
para("環境：本機開發伺服器連線正式 Supabase 資料庫（與正式站同一份資料），admin 測試帳號，Indian 品牌範圍。"
     "Playwright 自動化逐頁操作。結果如下：", size=10)
table(
    ["項目", "驗證內容", "結果"],
    [
        ["B-22", "Tab B 無 debug 文字、保留稽核日誌說明", ("✅ 通過", "EAF3DE")],
        ["B-21", "三組快篩按鈕存在；今日我的工單帶日期＋SA", ("✅ 通過", "EAF3DE")],
        ["B-19", "待派工橫幅實機抓到「12 張新工單待派工」", ("✅ 通過", "EAF3DE")],
        ["B-24", "拒絕原因圓餅圖渲染、SA 轉化率看板存在", ("✅ 通過", "EAF3DE")],
        ["B-23", "拒絕→五原因齊全、必填擋送出、選後啟用", ("✅ 通過", "EAF3DE")],
    ],
    widths=[0.8, 4.2, 1.1],
)
para("")
para("型別檢查（tsc）0 錯、ESLint 0 錯。13 項自動化斷言全數通過。", bold=True, color=GREEN)
para("補充：B-21 的 Walk-in 標籤與 B-24 的彩色圓餅扇形，需有對應資料（臨時進廠工單／新結構化拒絕）才會在畫面顯示；"
     "目前 Walk-in 標籤 0 個、圓餅圖顯示既有 2 筆未分類拒絕，均屬資料面正常狀態，功能本身已就緒。", size=9.5, color=GREY)

# ───────── 六、目錄結構規範 v3.0 ─────────
h1("六、目錄結構規範 v3.0 — 目錄調整完成")
para("收到《目錄結構規範 v3.0》後，我方先以正式 nav_nodes 實際資料逐項核對「規範要求」與「系統現況」是否相符，"
     "再調整。核對結論：規範的兩大訴求（分類註記消失、頁面名稱帶括號）對照我們的真實資料皆屬實，已全數處理。", size=10.5)

h2("6.1　核對結果（規範 vs 我們的真實 nav_nodes）")
table(
    ["規範要求", "我們現況核對", "處置"],
    [
        ["移除頁面名稱括號（6.1）", ("屬實：售後修護有 11 筆「（Step N）」名稱", "FDF3E3"), "已清為純功能名"],
        ["分類註記必須顯示（6.2）", ("屬實：資料在 DB（leaf section_group）但前端 renderer 未讀取 → 畫面不顯示", "FDF3E3"), "已修前端，全站分類註記點亮"],
        ["B-14~B-17 合併後移除轉RO", ("「預檢單轉RO（Step 3）」仍 active", "FDF3E3"), "已停用（is_active=false）"],
        ["四原則：導覽頁移除/設定收合/單一入口/品牌可控", ("大致符合", "EAF3DE"), "本批未違反"],
    ],
    widths=[2.3, 3.0, 1.0],
)
para("關鍵發現：分類註記「消失」的根因不在資料，而在前端——我們的 sidebar renderer 在三層樹模式下"
     "略過了 leaf 的 section_group 欄位，導致 CRM 等模組明明已設好「每日工作／追蹤管理／設定」卻不顯示。"
     "因此這不是純資料調整，需搭配一個前端小修。", size=10, color=AMBER)

h2("6.2　實際調整內容（僅 Indian 品牌）")
para("前端（2 檔，影響全站 sidebar）：", size=10, bold=True)
bullet("nav loader：三層樹模式改為讀取 leaf 的 section_group，作為分類註記帶入。")
bullet("SectionedTree 元件：同一群組內、功能頁之間，當分類改變時插入灰色斜體小標「── xxx ──」（不可點選，純視覺）。")
para("此修一次點亮全站既有分類註記——CRM 的「每日工作／追蹤管理／設定」、主管監看區的「每日監看」等也同步恢復顯示。", size=9.5, color=GREY)
para("資料（nav_nodes，Indian，共 21 筆）：", size=10, bold=True)
bullet("售後修護 SA 工單流程：10 筆移除「（Step N）」括號 + 設分類註記「Step 1–5 / Step 6–10」。")
bullet("售後修護 主管工作台：9 筆設分類註記「每日監看 / 審批與設定」。")
bullet("停用「預檢單轉RO（Step 3）」（B-14~B-17 合併後不再需要）。")
bullet("「工單查詢」入口改指向 /parts/aftersales/ro-search，讓 B-21 今日工單快篩有 sidebar 入口（原本指向工單列表頁、ro-search 為關閉狀態）。")

h2("6.3　調整後實機畫面（售後修護模組）")
add_shot("NAV-aftersales", "圖 v3.0：售後修護 sidebar — 分類註記已顯示、頁面名稱無括號，符合規範 §五", width=2.6)
para("CRM 模組的「每日工作／追蹤管理／設定」分類註記亦同步恢復顯示（同一前端修正生效）。", size=9.5, color=GREY)

h2("6.4　尚待確認 / 後續")
bullet("Ducati 品牌：本次僅調 Indian（海德生實際使用）。Ducati nav 結構未盤，待 Indian 驗收 OK 再比照套用。")
bullet("規範 6.3 的「href 指向 v2/v3」屬靜態 HTML 世界的版本概念；我方為 React 實作，對應路由本即正確，無需逐頁改 href（工單查詢入口已於本次一併修正）。")
bullet("其他模組若仍有括號名稱（如『盤點處理（條碼掃描）』等帶說明性括號），未在本批一併清除，待 Russell 確認是否要一律移除。")

para("")
para("— 報告結束 —", size=9, color=GREY)
doc.add_paragraph().alignment = WD_ALIGN_PARAGRAPH.CENTER

os.makedirs(OUTDIR, exist_ok=True)
doc.save(OUT)
print("saved:", OUT)

#!/usr/bin/env python3
# 實測回填：② 客服管理（CRM）模組場景清單對「真 React app」重測結果寫回 v1 docx（另存新檔保留原檔）
# 更新「支撐狀況」(col4) + 在「建議對策」(col6) 末追加【實測結果】+【本機 Playwright】
# 來源：subagent 對 localhost:3000 真 app 逐場景重測（docs/20260601/_subagent-crm-results.md）
#       + 本輪修正 CRM07-01/02（店長報表去 demo、換真 query）。verify：scripts/verify-crm-scenarios.mjs 13/13 綠。
from docx import Document
from docx.shared import RGBColor

SRC = "docs/20260601/03_DealerOS_場景驗證清單_客服管理模組_v1.docx"
OUT = "docs/20260601/03_DealerOS_場景驗證清單_客服管理模組_v1_實測回填.docx"

V = {
 "CRM01A-01": ("✅", "銷售客戶基盤真資料：sales-customer-base.ts 真連 customers（indian 52 筆）+ 姓名/電話/HABC/跟進狀態篩選 + 分頁 DataGrid，非寫死 8 筆假客戶。", True),
 "CRM01A-02": ("⚠️", "手卡 onSave 真呼叫 crm-sync upsertCustomer360 寫 customers.metadata.last_handcard_snapshot；但只 merge『已存在』客戶，查無時不會 insert 新客戶 → 新潛客手卡存了不會自動出現在基盤清單。差『查無則新增』一段。", True),
 "CRM01A-03": ("⚠️", "customers 為 sales/aftersales 共用同一張表（資料天生『都看得到』），但無真正的『交車 deal 完成→事件→建/標記售後客戶檔』觸發鏈（與 RS05-02 連動）。", True),
 "CRM02A-01": ("✅", "銷售電訪問卷 survey-templates 完整 CRUD + 版本管理 + SA 話術 + 啟用 toggle 真接 DB；建電訪任務時 getCallTaskLookups 撈 active 問卷餵 CRM03A 工作台，串通。", True),
 "CRM03A-01": ("⚠️", "call_tasks 表 + scheduled_at 真存在、工作台 getCallTaskBoardData 真讀（138 筆）；但無 cron/after() 在手卡儲存後自動產生 D+3/D+7 任務，現靠人工/seed 建。", True),
 "CRM03A-02": ("✅", "電訪結果真寫 DB：recordCallResultAction 真 update call_tasks，answered 時 insert nps_responses（連 call_task_id）；逾期由 scheduled_at runtime 計算示警。", True),
 "CRM04A-01": ("⚠️", "sales-dormant-leads 真讀 dormancy_status + 30-60/60-90/90+ 分桶（天數 runtime 真算）；但 dormancy_status 是 typed 靜態欄位，無 cron 把 active 自動降級成休眠。", True),
 "CRM04A-02": ("✅", "戰敗原因/競品分析真統計：getDormantLeadStats group lost_reason/competitor（lost 15 筆）；標戰敗走 sales-dormant-leads-actions 真寫 lost_reason/lost_at。", True),
 "CRM05A-01": ("✅", "銷售 NPS 真即時計算：sales-nps getSalesNpsDashboard 真讀 nps_responses（sales 38 筆）算 promoter-detractor + 門店分組 + 月度趨勢，非 +62 寫死。", True),
 "CRM06A-01": ("⚠️", "推播模板/活動/自動化規則 CRUD 真接 DB，createCampaignAction 真 insert push_campaigns（17 筆）；但不真發 LINE/SMS（無 notifications.dispatch），部分成效 KPI 為 demo 近似。專案已有 Notification Hub（feedback_ticket.created 已串），接此條為現成路。", True),
 "CRM01B-01": ("⚠️", "售後客戶基盤真 join work_orders 算最後回廠日/累計消費（33 工單，runtime 即時算，顯示效果≈同步）；但非『工單關閉事件→寫回客戶欄位』的自動觸發鏈。", True),
 "CRM01B-02": ("⚠️", "Desmo/保固/保險到期快篩按鈕在、真資料來自 customer_vehicles 主檔；自動到期推播見 CRM06B-01。", True),
 "CRM03B-01": ("⚠️", "售後電訪工作台 reuse CallTasksBoard（kind 鎖 aftersales）真讀 call_tasks + recordCallResult 真寫 nps（aftersales 104 筆）；但『工單關閉 D+3 滿意度自動產生 task』無觸發器。", True),
 "CRM02B-01": ("✅", "售後問卷與銷售問卷真隔離：reuse SurveyTemplatesBoard 但 kind 鎖死 'aftersales'，survey_templates 同表不同 kind（12 筆含兩 kind）獨立維護、互不影響。", True),
 "CRM04B-01": ("⚠️", "售後流失 crm-aftersales-dormant 真讀 aftersales_dormancy_status（dormant_60=5/120=4/180=4/lost=11 真分布）+ work_orders 算逾期天數（runtime 真）；但 dormancy_status typed 靜態欄位，無 cron 自動降級。", True),
 "CRM05B-01": ("✅", "售後 NPS 獨立真算：crm-aftersales-nps getAftersalesNpsDashboard 真讀 nps_responses kind=aftersales（104 筆）；store-overview 同源做 RS vs SA NPS 對比。", True),
 "CRM06B-01": ("⚠️", "售後推播自動化規則 reuse push 看板真讀 push_campaigns + toggleAutomationRuleAction 真寫 is_active；但自動推播規則可開關卻不真發送（無 cron 掃到期→dispatch），與 CRM06A-01 共用發送基礎待接。", True),
 "CRM07-01": ("✅", "【本輪修正】店長報表 store-overview getStoreOverview 多表真彙整（nps/sales_leads/call_tasks/work_orders/customers/sales_orders/push_campaigns）；本輪把殘留 demo 全換真 query：試駕轉化率改真值（完成 5→成交 3=60%）、RS 業績排行改 sales_orders 真聚合，算不出的個人指標（個人 NPS/D+3）誠實顯『—』或移除、不造假。tsc/eslint/天條/render 全綠。", True),
 "CRM07-02": ("✅", "【本輪修正】智能異常偵測由 getStoreOverview alerts[] 規則引擎真產生（NPS 批評者/逾期等規則，runtime 帶最近批評者客戶名）；本輪移除寫死的『王建宏 28%』示意 alert，改為末位 RS 落差過大才示警、帶真實姓名/台數。", True),
}

PW_NOTE = "【本機 Playwright 已驗】scripts/verify-crm-scenarios.mjs（admin+indian scope，13/13 路由 live render 全綠、無 error overlay）"

def cell_text(c):
    return "\n".join(p.text for p in c.paragraphs).strip()

doc = Document(SRC)
hit = pw = 0
for tbl in doc.tables:
    for row in tbl.rows:
        cells = row.cells
        if len(cells) < 7:
            continue
        sid = cell_text(cells[1]).strip()
        if sid not in V:
            continue
        status, result, tested = V[sid]
        cells[4].paragraphs[0].text = status
        p1 = cells[6].add_paragraph()
        r1 = p1.add_run(f"── 實測結果 2026-06-03：{result}")
        r1.font.color.rgb = RGBColor(0x18, 0x5F, 0xA5)
        if tested:
            p2 = cells[6].add_paragraph()
            r2 = p2.add_run(PW_NOTE)
            r2.bold = True
            r2.font.color.rgb = RGBColor(0x3B, 0x6D, 0x11)
            pw += 1
        hit += 1

n_ok = sum(1 for s, _, _ in V.values() if s == "✅")
n_warn = sum(1 for s, _, _ in V.values() if s == "⚠️")
n_no = sum(1 for s, _, _ in V.values() if s == "❌")

head = doc.paragraphs[0].insert_paragraph_before(
    "【實測回填 2026-06-03】本檔原評估係對舊靜態 Stitch HTML（CRM01A_v2.html 等、內含寫死假資料如 8 筆固定客戶、+62 NPS）所做，"
    "故幾乎全標 ❌/⚠️。本次改對「現在運行的真 React app」（localhost:3000、admin+indian scope）逐場景重測：先讀 code（page+元件+@/domain helper+server action）、SQL 確認 indian 真資料量，"
    f"再跑 Playwright live render。實況大幅好轉——全 {hit} 個場景對應 13 路由皆已 React 化 + 接 Supabase 真資料（customers 52／call_tasks 138／nps 142／leads 43…）、render 13/13 全綠、天條 0 違規、無任何 ❌。"
    f"重評：✅ {n_ok}（真頁+真資料+關鍵動作接 DB）／⚠️ {n_warn}（頁面與單頁讀寫已通、卡在跨模組『自動觸發鏈』未串）／❌ {n_no}。"
    "本輪另修正 CRM07-01/02（店長報表去除殘留 demo、全換真 query）。⚠️ 的共通根因是「跨模組自動觸發」尚未接通——"
    "缺口集中在：手卡查無則新增、交車自動建售後檔、D+3/D+7 與工單關閉後自動排程、休眠自動降級 cron、推播真發 LINE/SMS（Notification Hub 為現成路）。"
)
for r in head.runs:
    r.bold = True
    r.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)

doc.save(OUT)
print(f"回填 {hit} 場景（✅{n_ok}/⚠️{n_warn}/❌{n_no}，Playwright render 驗 {pw}）→ {OUT}")

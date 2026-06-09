# -*- coding: utf-8 -*-
"""C-28 FK橋接交接報告生成器"""
import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT = "docs/20260609/DealerOS_C28_FK橋接交接報告_2026-06-09.docx"
os.makedirs(os.path.dirname(OUT), exist_ok=True)

doc = Document()
style = doc.styles["Normal"]
style.font.name = "Microsoft JhengHei"
style.font.size = Pt(10.5)
style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft JhengHei")

NAVY   = RGBColor(0x1A, 0x3A, 0x5C)
GREEN  = RGBColor(0x3B, 0x6D, 0x11)
AMBER  = RGBColor(0x85, 0x4F, 0x0B)
RED    = RGBColor(0xCC, 0x00, 0x00)
GREY   = RGBColor(0x5A, 0x59, 0x55)
WHITE  = RGBColor(0xFF, 0xFF, 0xFF)

def shade_cell(cell, hexcolor):
    tcPr = cell._tc.get_or_add_tcPr()
    sh = OxmlElement("w:shd")
    sh.set(qn("w:val"), "clear")
    sh.set(qn("w:fill"), hexcolor)
    tcPr.append(sh)

def set_cell(cell, text, bold=False, color=None, size=9.5, fill=None, italic=False):
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.space_before = Pt(1)
    for i, line in enumerate(str(text).split("\n")):
        if i:
            p = cell.add_paragraph()
        r = p.add_run(line)
        r.bold = bold
        r.italic = italic
        r.font.size = Pt(size)
        r.font.name = "Microsoft JhengHei"
        r._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft JhengHei")
        if color:
            r.font.color.rgb = color
    if fill:
        shade_cell(cell, fill)

def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text)
    r.bold = True
    r.font.name = "Microsoft JhengHei"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft JhengHei")
    if level == 1:
        r.font.size = Pt(14)
        r.font.color.rgb = NAVY
    else:
        r.font.size = Pt(11.5)
        r.font.color.rgb = NAVY
    return p

def add_para(doc, text, color=None, size=10.5, indent=False, bold=False, space_before=2):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.space_before = Pt(space_before)
    if indent:
        p.paragraph_format.left_indent = Inches(0.25)
    for i, line in enumerate(text.split("\n")):
        if i:
            p.add_run().add_break()
        r = p.add_run(line)
        r.bold = bold
        r.font.size = Pt(size)
        r.font.name = "Microsoft JhengHei"
        r._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft JhengHei")
        if color:
            r.font.color.rgb = color
    return p

def add_code_block(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.left_indent = Inches(0.25)
    r = p.add_run(text)
    r.font.name = "Courier New"
    r.font.size = Pt(8.5)
    r.font.color.rgb = RGBColor(0x1A, 0x3A, 0x5C)

# ─── 標題頁 ──────────────────────────────────────────────────────────────────
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(6)
p.paragraph_format.space_after = Pt(2)
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("DealerOS")
r.bold = True; r.font.size = Pt(20); r.font.color.rgb = NAVY
r.font.name = "Microsoft JhengHei"
r._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft JhengHei")

p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
p2.paragraph_format.space_after = Pt(2)
r2 = p2.add_run("C-28 FK橋接交接報告")
r2.bold = True; r2.font.size = Pt(15); r2.font.color.rgb = NAVY
r2.font.name = "Microsoft JhengHei"
r2._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft JhengHei")

p3 = doc.add_paragraph()
p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
p3.paragraph_format.space_after = Pt(12)
r3 = p3.add_run("work_orders ↔ repair_orders FK 橋接 + 增項零件實體出庫\n2026-06-09　｜　Russell Hung × Claude Sonnet 4.6")
r3.font.size = Pt(10); r3.font.color.rgb = GREY
r3.font.name = "Microsoft JhengHei"
r3._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft JhengHei")

doc.add_paragraph()

# ─── 一、任務背景 ─────────────────────────────────────────────────────────────
add_heading(doc, "一、任務背景")
add_para(doc, "DealerOS 存在兩套獨立的工單系統：")
add_para(doc, "● work_orders（後台 ERP 工單）：ro_no 格式 RO-DUC-XXXXX，紀錄維修作業的技術內容、工時與零件用量，由後台人員操作。", indent=True)
add_para(doc, "● repair_orders（前台服務工單）：ro_code 格式 RP-CP-YYMMDD-NNN，紀錄客戶接待、預約、SA 派工與關單流程，由前台服務顧問操作。", indent=True)
add_para(doc, "")
add_para(doc, "C-28 問題：兩張表在資料庫層完全沒有直接 FK 橋接，僅能透過共同的 customer_id / vehicle_id 做業務語意層隱性關聯，不構成結構性橋接。此外，stock_issues.ro_id 雖在 code 層被當作 work_orders.id 的軟參照，但缺乏正式 FK constraint，存在資料一致性風險。")
add_para(doc, "")
add_para(doc, "C-28 任務目標：在 DB 層建立 work_orders ↔ repair_orders 的正式 FK 橋接，並將 stock_issues.ro_id 補上 FK constraint，確保增項出庫單能追溯到對應的 ERP 工單。")

# ─── 二、技術方案 ─────────────────────────────────────────────────────────────
add_heading(doc, "二、技術方案")
add_para(doc, "採用「work_orders 側新增 repair_order_id 欄位」的單向橋接策略：")
add_para(doc, "① repair_orders 已有 pre_inspection_id，若需雙向查詢透過 work_orders.repair_order_id 反查即可，不需在 repair_orders 側加冗餘欄位。", indent=True)
add_para(doc, "② 業務語意：後台 WO 引用前台 RO，依賴方向正確。", indent=True)
add_para(doc, "③ ON DELETE SET NULL：當 repair_order 被刪除時，work_orders 不受影響，橋接欄位歸 null（防禦性設計）。", indent=True)
add_para(doc, "")
add_para(doc, "自動閉環機制：transferToROAction（RO Handoff 開單）成功建立 repair_order 後，step 3b 自動透過 appointment_id 回填 work_orders.repair_order_id。雙重防護：同 brand_id + 同 appointment_id + repair_order_id IS NULL，避免重複回填；找不到對應 work_order 不中斷主流程（fire-and-forget）。")

# ─── 三、實施內容 ─────────────────────────────────────────────────────────────
add_heading(doc, "三、實施內容")

add_heading(doc, "3-1　DB Migration（已 apply 至 Supabase Production）", level=2)
add_code_block(doc, "ALTER TABLE work_orders")
add_code_block(doc, "  ADD COLUMN IF NOT EXISTS repair_order_id uuid")
add_code_block(doc, "  REFERENCES repair_orders(id) ON DELETE SET NULL;")
add_code_block(doc, "")
add_code_block(doc, "CREATE INDEX IF NOT EXISTS idx_work_orders_repair_order_id")
add_code_block(doc, "  ON work_orders(repair_order_id) WHERE repair_order_id IS NOT NULL;")
add_code_block(doc, "")
add_code_block(doc, "ALTER TABLE stock_issues")
add_code_block(doc, "  ADD CONSTRAINT stock_issues_ro_id_fkey")
add_code_block(doc, "  FOREIGN KEY (ro_id) REFERENCES work_orders(id) ON DELETE SET NULL;")

add_heading(doc, "3-2　程式碼變更", level=2)

t = doc.add_table(rows=1, cols=2)
t.style = "Table Grid"
t.columns[0].width = Inches(2.2)
t.columns[1].width = Inches(4.0)
hdr = t.rows[0].cells
set_cell(hdr[0], "檔案路徑", bold=True, fill="1A3A5C", color=WHITE, size=9.5)
set_cell(hdr[1], "改動說明", bold=True, fill="1A3A5C", color=WHITE, size=9.5)

rows_data = [
    (
        "src/lib/aftersales/ro-handoff-actions.ts",
        "transferToROAction 成功建立 repair_order 後，新增 step 3b：若 PI 帶有 appointment_id，對 work_orders 回填 repair_order_id（條件：同 brand_id + 同 appointment_id + repair_order_id IS NULL），完成 FK bridge 閉環。找不到對應 work_order 不影響主流程。"
    ),
    (
        "src/lib/database.types.ts",
        "重新從 Supabase 產生 TypeScript 型別（generate_typescript_types）。work_orders Row/Insert/Update 現在包含 repair_order_id: string | null；stock_issues FK constraint 也已反映在 schema 中。"
    ),
]
for path, desc in rows_data:
    row = t.add_row().cells
    set_cell(row[0], path, size=8.5, color=NAVY)
    set_cell(row[1], desc, size=9.5)

doc.add_paragraph()

# ─── 四、驗收結果 ─────────────────────────────────────────────────────────────
add_heading(doc, "四、驗收結果")

vt = doc.add_table(rows=1, cols=3)
vt.style = "Table Grid"
vt.columns[0].width = Inches(2.4)
vt.columns[1].width = Inches(1.2)
vt.columns[2].width = Inches(2.6)
vh = vt.rows[0].cells
set_cell(vh[0], "驗收項目", bold=True, fill="1A3A5C", color=WHITE, size=9.5)
set_cell(vh[1], "結果", bold=True, fill="1A3A5C", color=WHITE, size=9.5)
set_cell(vh[2], "說明", bold=True, fill="1A3A5C", color=WHITE, size=9.5)

verify_rows = [
    ("DB 欄位（work_orders.repair_order_id）", "✅ PASS", "type=uuid, nullable=true，已確認存在"),
    ("FK Constraint", "✅ PASS", "work_orders_repair_order_id_fkey（ON DELETE SET NULL）及 stock_issues_ro_id_fkey 均已生效"),
    ("Supabase 天條（UI 禁直連）", "✅ 0 violations", "4 筆 grep hits 均為既有 layout.tsx 問題或程式碼註解，非本次改動"),
    ("TypeScript 編譯", "✅ 0 errors", "npx tsc --noEmit 乾淨通過"),
    ("Domain logic 審查", "✅ PASS", "step 3b fire-and-forget 邏輯合規，src/lib/ 層直接 import supabase/server 合法"),
]
for check, status, detail in verify_rows:
    vrow = vt.add_row().cells
    set_cell(vrow[0], check, size=9.5)
    set_cell(vrow[1], status, size=9.5, bold=True, color=GREEN)
    set_cell(vrow[2], detail, size=9.0)

doc.add_paragraph()

# ─── 五、下一步（未來工作） ───────────────────────────────────────────────────
add_heading(doc, "五、下一步（未來工作）")
add_para(doc, "C-28 DB 橋接已完成，以下 UI 層補足工作建議後續 session 接手：")

next_steps = [
    ("work-orders.ts domain helper 擴充",
     "新增 WorkOrderDetail type（含 repair_order_id？）及 getWorkOrderWithRepairOrder(id) helper，join repair_orders 取 ro_code 供 UI 顯示關聯服務單號。"),
    ("workorder-form-types.ts + workorder-actions.ts",
     "WorkOrderFormValues 加 repair_order_id?: string | null；createWorkOrderAction / updateWorkOrderAction 支援手動指定橋接。"),
    ("issues.ts 增項出庫閉環",
     "pickForWorkOrder() 建單時從 work_order 取出 repair_order_id 寫入 stock_issues.metadata；listIssues 補 repair_order_code 欄位供 UI 顯示。"),
    ("repair-order-actions.ts 關單路徑補強",
     "confirmRepairOrderAction 關單時若對應 work_order 存在，自動回寫 work_orders.repair_order_id（目前只有 transferToROAction 開單時回填）。"),
    ("layout.tsx 既有天條問題",
     "grep 查到 layout.tsx 直連 supabase，屬 C-28 之前既有問題，建議另立 ticket 修正。"),
]
for i, (title, detail) in enumerate(next_steps, 1):
    add_para(doc, f"{i}. {title}", bold=True, size=10.5, space_before=6)
    add_para(doc, detail, indent=True, size=10.0, color=GREY)

doc.add_paragraph()

# ─── 六、Git Commit 建議 ──────────────────────────────────────────────────────
add_heading(doc, "六、Git Commit 建議")
add_code_block(doc, "feat(c28): work_orders ↔ repair_orders FK 橋接，增項出庫帶 repair_order_id")
add_code_block(doc, "")
add_code_block(doc, "- ALTER TABLE work_orders ADD COLUMN repair_order_id uuid REFERENCES repair_orders(id) ON DELETE SET NULL")
add_code_block(doc, "- ALTER TABLE stock_issues ADD CONSTRAINT stock_issues_ro_id_fkey FOREIGN KEY (ro_id) REFERENCES work_orders(id)")
add_code_block(doc, "- ro-handoff-actions.ts transferToROAction step 3b：開單後自動回填 work_orders.repair_order_id")
add_code_block(doc, "- database.types.ts 重新產 TypeScript 型別")

doc.add_paragraph()
add_para(doc, "─" * 60, color=GREY, size=9)
add_para(doc, "DealerOS C-28 交接報告　｜　2026-06-09　｜　機密文件", color=GREY, size=9)

doc.save(OUT)
print(f"✅ 報告已儲存：{OUT}")

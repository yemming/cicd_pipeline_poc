# -*- coding: utf-8 -*-
"""產出「給人看」的補做完成報告 Word 檔（含截圖）。"""
import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

BASE = os.path.dirname(os.path.abspath(__file__))
SHOTS = os.path.join(BASE, "shots")
OUT = os.path.join(BASE, "補做完成報告_退料閉環TL工單_20260617.docx")

NAVY = RGBColor(0x1A, 0x3A, 0x5C)
GREEN = RGBColor(0x0F, 0x6E, 0x56)
RED = RGBColor(0xCC, 0x00, 0x00)
GREY = RGBColor(0x5A, 0x59, 0x55)
FONT = "Microsoft JhengHei"

doc = Document()
style = doc.styles["Normal"]
style.font.name = FONT
style.font.size = Pt(10.5)
try:
    style.element.rPr.rFonts.set(
        "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}eastAsia", FONT)
except Exception:
    pass


def set_font(run, size=10.5, bold=False, color=None):
    run.font.name = FONT
    run.font.size = Pt(size)
    run.font.bold = bold
    if color is not None:
        run.font.color.rgb = color
    try:
        run._element.rPr.rFonts.set(
            "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}eastAsia", FONT)
    except Exception:
        pass


def h1(text):
    p = doc.add_paragraph()
    p.space_before = Pt(10)
    r = p.add_run(text)
    set_font(r, 16, True, NAVY)
    return p


def h2(text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    set_font(r, 13, True, NAVY)
    return p


def para(text, size=10.5, color=None, bold=False, space_after=4):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    r = p.add_run(text)
    set_font(r, size, bold, color)
    return p


def bullet(text, color=None):
    p = doc.add_paragraph(style="List Bullet")
    r = p.add_run(text)
    set_font(r, 10.5, False, color)
    return p


def caption(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    set_font(r, 9, False, GREY)
    p.paragraph_format.space_after = Pt(10)


def add_image(filename, width_in=6.3, cap=None):
    path = os.path.join(SHOTS, filename)
    if os.path.exists(path):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(path, width=Inches(width_in))
        if cap:
            caption(cap)
    else:
        para(f"[缺圖：{filename}]", color=RED)


def table(headers, rows):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Light Grid Accent 1"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, hd in enumerate(headers):
        c = t.rows[0].cells[i]
        c.text = ""
        r = c.paragraphs[0].add_run(hd)
        set_font(r, 10, True)
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ""
            r = cells[i].paragraphs[0].add_run(str(val))
            set_font(r, 9.5)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return t


# ── 封面 ──
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.LEFT
r = title.add_run("補做完成報告")
set_font(r, 22, True, NAVY)
sub = doc.add_paragraph()
r = sub.add_run("退料閉環與 TL 借用測試工單 — 兩項必須補做，一次性回覆")
set_font(r, 12, False, GREY)

meta = doc.add_paragraph()
for t_, b_ in [("回覆對象：", "Russell Hung（DealerOS）"),
               ("　｜　日期：", "2026-06-17")]:
    set_font(meta.add_run(t_), 10, False, GREY)
    set_font(meta.add_run(b_), 10, True)

doc.add_paragraph()

# ── 摘要 ──
h1("摘要（先講結論）")
table(["項目", "狀態", "證據"], [
    ["Item 1 — TL 走正式 repair-pick 倉管領料", "✅ 完成", "倉管在 repair-pick 看到 TL 工單、可預覽配置過帳（3 張截圖）"],
    ["Item 2 — prefix constraint migration 進版控", "✅ 完成", "supabase/migrations 補 3 檔，與正式站 DB 一致"],
    ["澄清 — stock_item 中間狀態", "✅ 已說明", "見第 3 節"],
    ["澄清 — demo 資料保留", "✅ 已記錄", "兩筆 demo 工單未動"],
])
para("我方接受批評。「非閉環關鍵、可跳過」的判斷不該由執行端自行做出；遇到技術困難"
     "（work_orders 橋接）正確做法是回報讓貴方決定。本次兩項補做已落地，倉管已重新成為"
     "零件進出的必經關卡。", space_after=6)
para("部署狀態：本輪程式變更尚未上線；demo 橋接資料已寫入正式站資料庫，故現有部署的 "
     "repair-pick 已可看到 TL 工單（截圖即在正式站擷取）。自動橋接的程式（未來 TL 工單建立即"
     "自動進倉管領料清單）將於下次部署生效。", color=GREY, space_after=8)

# ── Item 1 ──
h1("Item 1：TL 借料走正式 repair-pick 倉管領料")
h2("為什麼原本走不通")
para("倉管的「維修領料（repair-pick）」清單與領料預覽，完全以維修工單（work_orders）為來源，"
     "不認 TL 那一張 repair_orders。TL 借用測試工單建立時不會產生對應的 work_orders，"
     "所以永遠不出現在倉管的領料清單——上一版才走了結案頁自行逐行處置的捷徑。")
h2("補做的橋接")
bullet("TL 工單建立 / 借料明細變動時，自動「橋接」出一筆對應的倉管維修工單，並把借出零件帶過去。")
bullet("TL 工單因此自動進入倉管「待領料工單」清單；倉管選工單 → 系統算出庫配置 → 倉管確認過帳。")
bullet("零件出庫從此必經倉管之手，每次領料都有領料單（GI 單）可稽核，責任歸屬清楚。")
para("配套一筆資料庫調整：TL 是內部借用、沒有客戶，原本工單表要求一定要有客戶才能建立，"
     "已放寬此限制（並補 migration 檔，見 Item 2）。落地前已確認不影響權限隔離、搜尋索引、既有資料。",
     color=GREY)

h2("驗證截圖（正式站 · Indian）")
add_image("A_倉管_待領料工單清單含TL.png", 6.3,
          "圖 A：倉管「維修領料」→「待領料工單」清單，已包含 TL-IN-260616-901（料件 3 項/3 件，附『備料』鈕）。")
add_image("B_倉管_新增領料可選TL工單.png", 6.3,
          "圖 B：倉管新增領料「綁定 RO 工單」，TL-IN-260616-901 可被勾選領料（狀態 dispatched、料件 3 項）。")
add_image("C_倉管_TL工單FIFO配置可過帳.png", 6.3,
          "圖 C：選 TL 工單後的出庫配置預覽（可過帳）— 3 筆零件各需求 1、可用 26，合計 NT$750，倉管按『一鍵領料並過帳』即正式發料。")
para("註：截圖只到「預覽可過帳」為止，未實際點過帳，以免動到貴方要保留的 demo 庫存。",
     color=GREY)

# ── Item 2 ──
h1("Item 2：DB constraint 修改的 migration 記錄")
para("釐清：本專案早期資料庫變更多以工具直接套用到雲端，雲端的 migration 紀錄表完整保存了"
     "version 與 SQL（單一事實來源），但未自動在程式碼倉庫落地成檔。"
     "prefix 白名單加 TL 的那筆 ALTER 即屬此情況——貴方的批評成立。")
para("補做：supabase/migrations/ 補上 3 個 migration 檔（內容與正式站逐字一致），"
     "並更新規範文件，明訂「所有涉及流程控制、行為稽核、資料結構的 DB 變更，一律要有 migration 檔進版控」。",
     space_after=4)
table(["檔案", "說明"], [
    ["20260616132008_create_parts_return_requests.sql", "退料閉環核心表（同輪補檔）"],
    ["20260616141106_add_tl_to_repair_orders_prefix_p1.sql", "貴方點名的 prefix 白名單加 TL"],
    ["20260617033103_tl_bridge_work_orders_customer_id_nullable.sql", "本輪 customer_id 放寬（橋接需求）"],
])
para("一致性確認（2026-06-17 查正式站）：prefix 約束已含 TL、work_orders.customer_id 已可為空、"
     "3 筆 version 皆在雲端 migration 表有紀錄，三者與 codebase 檔案一致。", color=GREY)

# ── 澄清一 ──
h1("澄清一：取消到倉管確認之間，stock_item 的中間狀態")
para("走正式 repair-pick 後，TL 借料的庫存生命週期：倉管過帳 → 該批料件扣庫、標記為已出庫"
     "（issued），並留下領料單；TL 結案選「退回庫房」→ 建一筆『退料待確認』記錄（pending），"
     "此時庫存不立即回補；倉管在退料入庫實際點收確認 → 才 insert 一筆全新的可用庫存回補。")
para("直球回答：那段期間，已出庫的那批料件維持「已出庫（issued）」，同時有一筆『退料待確認』"
     "記錄作為「預期退回、待倉管點收」的可見證明。", bold=True)
para("我方認為這在物理上是相符的：料件確實還在外面（技師手上／退回途中），尚未實際回到貨架；"
     "報表此時顯示它不在可用庫存是正確的。帳面與實物的對齊點，是倉管實體點收那一刻，而不是"
     "結案決定那一刻——這正符合「零件進出必經倉管」的原則：回庫也要倉管點頭才算數。")
para("若貴方希望在庫存查詢另立一個明確的「退料在途」可視狀態，這屬展示層強化，我方可另開規格"
     "提案，等貴方裁示再做，不自行決定。", color=GREY)

# ── 澄清二 ──
h1("澄清二：測試資料殘留")
para("依指示，正式站兩筆 demo 工單暫時保留、不覆蓋、不刪除，待海德生 UAT 後再清除：")
table(["單號", "用途"], [
    ["RP-CP-260616-901", "退料閉環 demo（3 零件行 + 1 追加項目）"],
    ["TL-IN-260616-901", "TL 借用測試 demo（3 零件行 + 1 工項）"],
])
para("本次為 TL demo 新增的橋接工單（供截圖驗證，亦屬可保留 demo 資料）：work_order "
     "6b2ab934…（對應 TL-IN-260616-901，3 筆零件）。截圖驗證停在「預覽可過帳」、未實際過帳，"
     "故 demo 庫存未被改動。", color=GREY)

# ── 結語 ──
h1("最後")
para("我方清楚認識到：事關流程控制與行為稽核的邏輯不能草率；遇到技術困難，正確做法是回報讓"
     "貴方決定，而不是自行判斷哪個重要、哪個可跳過。本次補做已讓倉管重新回到零件進出的關卡，"
     "並把資料庫變更納入版本控制。")
para("待裁示：(1) 本輪程式變更的上線時點；(2) 庫存報表是否要新增「退料在途」可視狀態。",
     color=GREY)

doc.save(OUT)
print("WROTE", OUT)

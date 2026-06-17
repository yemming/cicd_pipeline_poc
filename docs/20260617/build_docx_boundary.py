# -*- coding: utf-8 -*-
"""裁示三回覆 Word 檔，含三狀態 closeup 截圖。"""
import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

BASE = os.path.dirname(os.path.abspath(__file__))
SHOTS = os.path.join(BASE, "shots-boundary")
OUT = os.path.join(BASE, "回覆_借料未還兩邊界驗證_20260617.docx")
NAVY = RGBColor(0x1A, 0x3A, 0x5C); GREY = RGBColor(0x5A, 0x59, 0x55); RED = RGBColor(0xCC, 0, 0)
FONT = "Microsoft JhengHei"; EA = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}eastAsia"

doc = Document(); s = doc.styles["Normal"]; s.font.name = FONT; s.font.size = Pt(10.5)
try: s.element.rPr.rFonts.set(EA, FONT)
except Exception: pass

def sf(r, sz=10.5, b=False, c=None):
    r.font.name = FONT; r.font.size = Pt(sz); r.font.bold = b
    if c is not None: r.font.color.rgb = c
    try: r._element.rPr.rFonts.set(EA, FONT)
    except Exception: pass
def h1(t):
    p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(10); sf(p.add_run(t), 16, True, NAVY)
def h2(t):
    p = doc.add_paragraph(); sf(p.add_run(t), 13, True, NAVY)
def para(t, sz=10.5, c=None, b=False, after=4):
    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(after); sf(p.add_run(t), sz, b, c)
def bullet(t, c=None):
    p = doc.add_paragraph(style="List Bullet"); sf(p.add_run(t), 10.5, False, c)
def cap(t):
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER; sf(p.add_run(t), 9, False, GREY); p.paragraph_format.space_after = Pt(10)
def img(fn, w=6.6, c=None):
    path = os.path.join(SHOTS, fn)
    if os.path.exists(path):
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER; p.add_run().add_picture(path, width=Inches(w))
        if c: cap(c)
    else: para(f"[缺圖 {fn}]", c=RED)
def table(headers, rows):
    t = doc.add_table(rows=1, cols=len(headers)); t.style = "Light Grid Accent 1"; t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, hd in enumerate(headers):
        cc = t.rows[0].cells[i]; cc.text = ""; sf(cc.paragraphs[0].add_run(hd), 10, True)
    for row in rows:
        cs = t.add_row().cells
        for i, v in enumerate(row):
            cs[i].text = ""; sf(cs[i].paragraphs[0].add_run(str(v)), 9.5)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)

p = doc.add_paragraph(); sf(p.add_run("回覆 — 借料未還兩項邊界情境驗證（裁示三）"), 19, True, NAVY)
p = doc.add_paragraph(); sf(p.add_run("致 Russell Hung（DealerOS）　｜　2026-06-17　｜　一次性回覆（裁示一～四）"), 11, False, GREY)
doc.add_paragraph()

h1("摘要")
table(["裁示", "回應"], [
    ["一 · 借料未還採三階", "✅ 已採三階（≤3天藍／4-7天amber／>7天紅⚠），門檻為可調常數"],
    ["二 · 暫不鋪 return-in／列表", "✅ 知悉；本輪僅 TL detail 一個 surface，擴展留待 UAT 回饋"],
    ["三 · 兩邊界須驗證到不模糊", "✅ 判定改逐件量化 + 真實 e2e 三狀態驗證（本文重點）"],
    ["四 · 測試資料清除", "✅ 知悉；本輪另建 3 筆測試 TL，一併列入清除清單"],
])
para("我方完全認同裁示三的標準與 B-01 連結：顯示邏輯本身不能把警示蓋掉。一個在常見情境顯示含糊或"
     "誤導的提醒，會讓倉管學會不信任、不再多看一眼，等同回到「沒有倉管紀錄的進出沒人負責」的原始風險。"
     "因此判定邏輯已從粗略的工單層級，改成逐件量化。", after=8)

h1("核心修正：判定從「工單層級」改成「逐件量化」")
para("某零件「未還量」＝ 已出庫量 − 已確認回庫量(confirmed) − 已處置量(收費／吸收／轉單)。"
     "pending／逾期的退料『不』扣除——B-01：倉管實體點收前，零件還在外面。", after=4)
bullet("chip：「借料未還 N 項（最久 D 天）」，顏色由最久未還件決定（先讓人看到最該擔心那筆）。")
bullet("detail 逐件明細：每件列「已借 D 天（分級色）· 品名 ×數量 ·（出庫時間）」，並標「另有 X 筆退料待倉管點收」。")

h1("邊界一：分批出庫，天數起算是否誤導")
h2("結論：現行為「原子出庫」，分批不可達 → 不存在天數混算")
para("倉管領料流程查證：工單一旦領過料即從待領清單排除、新增領料清單前端硬過濾、URL 繞路無效、後端不支援"
     "「領剩餘」。故一張 TL 的零件只能一次領料、原子出庫，所有零件共用同一出庫時間——不可能出現「老件與"
     "新件被算成同一天數」的誤導，因為根本沒有分不同日期出庫這條路。")
para("但仍做到逐件清晰（不只給一個籠統數字）：detail 逐件明細把每件的出庫時間與天數分開列、顏色取最久件。")
img("D1_after-pick-2items.png", 6.6,
    "圖 1：TL-IN-260617-003 借 2 件、一次領料後 —— chip「借料未還 2 項（最久 0 天）」＋逐件明細，兩件各列「已借 0 天 · 出庫 16:53」，同一出庫時間逐件攤開，零模糊。日後若開放分批，明細會各自顯示不同出庫時間與天數。")

h1("邊界二：部分歸還，chip 是否準確（裁示三 MUST-FIX）")
para("真實 e2e（全程 UI、無 SQL 塞值）：同一張 TL（借 2 件）走三個狀態，驗證 chip 在「有退料動作」後是否仍"
     "忠實反映還在外面的零件。", after=4)
table(["狀態", "操作", "chip"], [
    ["① 領料後", "倉管過帳，2 件出庫", "借料未還 2 項（最久 0 天）"],
    ["② 結案+退料申請後", "tl-close 兩件都「退回庫房」、工單已關單", "借料未還 2 項（仍是！）"],
    ["③ 倉管確認 1 件回庫後", "return-in 確認 1 件", "借料未還 1 項"],
])
img("D2_after-close-still-2items.png", 6.6,
    "圖 2（MUST-FIX 關鍵）：工單已關單、已建 2 筆退料申請，但因倉管尚未實體點收，chip 仍顯示「借料未還 2 項」，明細標「另有 2 筆退料待倉管點收」。不會因為有退料動作就顯示沒事。")
img("D3_after-confirm-1item.png", 6.6,
    "圖 3：倉管確認 1 件回庫後，chip 準確降為「借料未還 1 項」，明細只剩未還的煞車#002。降的是「已確認回庫」那件，不是把警示整個關掉。")
h2("直球回答裁示三兩問")
para("1. chip 是否持續正確顯示借料未還？是。狀態②工單已關單＋已建退料申請，但倉管未點收前 chip 仍顯示"
     "「借料未還 2 項」；狀態③確認 1 件後準確降為 1 項。不會因有退料動作就顯示沒事。", b=False)
para("2. 天數是否被誤判降低？不會。三狀態天數一律「最久 0 天」（從出庫日算，不受退料動作影響）。outstanding 只看"
     "逐件未還量是否>0，pending 退料不扣未還量。只要還有任一件未經倉管實體點收回庫，chip 一定持續顯示「借料未還」、"
     "絕不出現「沒事」。")
para("為何如此設計（對齊 B-01）：「未還」的對齊點是倉管實體點收回庫那一刻，不是技師喊還、也不是工單結案。"
     "pending 退料＝申請退、倉管還沒點收＝零件物理上還在外面。回庫也要倉管點頭才算數，顯示邏輯不替倉管把警示提前關掉。",
     c=GREY, after=8)

h1("裁示一／二／四 回應")
bullet("一：三階（≤3天 資訊藍／4-7天 amber／>7天 紅⚠），門檻為常數（INFO_MAX=3、WARNING_MIN=7），要微調直接改。")
bullet("二：本輪僅 TL detail 一個 surface（含逐件明細）；return-in 列／工單列表擴展留待真人盲測回饋。")
bullet("四：本輪另建 3 筆測試 TL（TL-IN-260617-001/002/003，皆已實際出庫、部分退料/回補），併入清除清單，清除時連同庫存回補處理。")

h1("等裁示")
bullet("三階門檻（3／7 天）如需調整，告知即可（改常數）。")
bullet("是否進入真人盲測，以決定要不要把「借料未還」擴展到 return-in 列／工單列表。")

doc.save(OUT); print("WROTE", OUT)

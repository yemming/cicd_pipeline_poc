# -*- coding: utf-8 -*-
"""回覆（上線時間 + 借料未還）Word 檔，含真實 e2e 截圖。"""
import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

BASE = os.path.dirname(os.path.abspath(__file__))
SHOTS = os.path.join(BASE, "shots-real")
OUT = os.path.join(BASE, "回覆_上線時間與借料未還_20260617.docx")

NAVY = RGBColor(0x1A, 0x3A, 0x5C)
GREY = RGBColor(0x5A, 0x59, 0x55)
RED = RGBColor(0xCC, 0x00, 0x00)
FONT = "Microsoft JhengHei"

doc = Document()
st = doc.styles["Normal"]
st.font.name = FONT
st.font.size = Pt(10.5)
EA = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}eastAsia"
try:
    st.element.rPr.rFonts.set(EA, FONT)
except Exception:
    pass


def sf(run, size=10.5, bold=False, color=None):
    run.font.name = FONT
    run.font.size = Pt(size)
    run.font.bold = bold
    if color is not None:
        run.font.color.rgb = color
    try:
        run._element.rPr.rFonts.set(EA, FONT)
    except Exception:
        pass


def h1(t):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    sf(p.add_run(t), 16, True, NAVY)


def h2(t):
    p = doc.add_paragraph()
    sf(p.add_run(t), 13, True, NAVY)


def para(t, size=10.5, color=None, bold=False, after=4):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(after)
    sf(p.add_run(t), size, bold, color)


def bullet(t, color=None):
    p = doc.add_paragraph(style="List Bullet")
    sf(p.add_run(t), 10.5, False, color)


def cap(t):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sf(p.add_run(t), 9, False, GREY)
    p.paragraph_format.space_after = Pt(10)


def img(fn, w=6.3, c=None):
    path = os.path.join(SHOTS, fn)
    if os.path.exists(path):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(path, width=Inches(w))
        if c:
            cap(c)
    else:
        para(f"[缺圖 {fn}]", color=RED)


def table(headers, rows):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Light Grid Accent 1"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, hd in enumerate(headers):
        c = t.rows[0].cells[i]
        c.text = ""
        sf(c.paragraphs[0].add_run(hd), 10, True)
    for row in rows:
        cells = t.add_row().cells
        for i, v in enumerate(row):
            cells[i].text = ""
            sf(cells[i].paragraphs[0].add_run(str(v)), 9.5)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


# 封面
p = doc.add_paragraph()
sf(p.add_run("回覆 — 上線時間說明 + 「借料未還」可視狀態"), 20, True, NAVY)
p = doc.add_paragraph()
sf(p.add_run("致 Russell Hung（DealerOS）　｜　2026-06-17　｜　一次性回覆（4 點）"), 11, False, GREY)
doc.add_paragraph()

h1("摘要")
table(["要求", "狀態", "證據"], [
    ["① 自動橋接 push 確切日期 + 原因", "✅ 已上線", "2 commit、Zeabur RUNNING 時間戳"],
    ["② 全新 TL 工單真實 e2e（未塞資料）", "✅ 完成", "TL-IN-260617-001 全程 UI 截圖"],
    ["③「借料未還」門檻 + UI + 實作", "✅ 完成", "chip 已上線、真實工單截圖"],
    ["④ 一次性回覆", "✅ 本檔", "—"],
])
para("感謝對「澄清一」中間狀態論證的肯定。要求二的設計哲學——系統不該只忠實記錄狀態，"
     "要主動把風險翻譯成警語，讓對的人在對的時間看到——我方已內化，並以此實作「借料未還」。",
     after=8)

# 要求一
h1("要求一：自動橋接的 push 日期與原因")
table(["commit", "內容", "部署（Asia/Taipei）"], [
    ["10d9b22", "TL 橋接走正式 repair-pick + 3 個 migration 補檔", "12:07 Zeabur RUNNING"],
    ["d21d78a", "「借料未還」可視狀態 + 時間分級", "13:58 Zeabur RUNNING（含前一 commit）"],
])
para("兩者皆已上正式站，自動橋接 code 已生效。", after=4)
h2("為什麼前一份報告時還沒 push")
para("卡點是內部部署門檻：本專案正式站 push 需經 owner 對「上正式環境」放行——這是既有工程"
     "紀律（避免未經確認就動到收費營業環境），不是技術問題、也不是測試環境卡關。")
para("我方接受批評：前一份用「等 Ming 確認」一句帶過、沒講清楚等什麼、預計多久，確實無法被"
     "追蹤。正確寫法應是「卡點＝正式站部署的 owner 放行關卡；預計 owner 確認後當日 push」。"
     "owner 已於 6/17 當日放行，兩個 commit 皆已 push 並部署成功。日後涉及部署等待，一律寫明"
     "等什麼關卡、預計時點，把進度說明當稽核紀錄寫。", color=GREY, after=8)

# 要求二
h1("要求二：全新 TL 工單的真實 e2e（未經人工塞資料）")
para("上一份三張截圖是把 demo 橋接工單直接寫進正式站 DB，模擬「橋接生效後畫面長怎樣」——"
     "證明設計方向對，不是「真的開一張新 TL、系統自動觸發橋接」。本次補上真正的端到端驗證："
     "在正式站、走正式 UI，從零開一張全新 TL 工單，全程不碰資料庫（連簽名都是在 canvas 上實際"
     "畫的）。")
para("新工單：TL-IN-260617-001（林志玲 · Indian Chief Vintage · RDC-2201）。", bold=True)
img("05_repair-pick-pending-NEW-TL.png", 6.3,
    "圖 1（關鍵）：倉管「維修領料 → 待領料工單」自動出現剛建立的 TL-IN-260617-001（料件 1 項/2 件，傳動系統零件 #001）。此工單由正式 UI 即時建立，橋接 code 自動觸發，非人工塞資料。")
img("03_tl-ro-created.png", 6.3,
    "圖 2：TL 工單開立成功頁（SA 與技師簽名均為 Playwright 在 canvas 實際手繪）。")
img("07_pick-posted-NEW-TL.png", 6.3,
    "圖 3：倉管預覽 FIFO 配置後，一鍵領料並過帳 —— 零件正式出庫、寫領料單、扣庫。倉管全程參與發料確認。")
para("唯一前置 scaffold 是一筆乾淨預約（閘門頁規定 TL 由預約轉開，屬既有流程，非 TL 工單本身）。"
     "TL 工單與整條橋接、出庫、借料未還 chip，全部由正式 app 即時產生。驗證腳本 tests/e2e/tl-real-e2e.mjs。",
     color=GREY)
para("此真實工單 TL-IN-260617-001（id bdc7934a…）已實際出庫 2 件，請與兩筆 demo 一併列入測試資料"
     "殘留清單，清除時需連同回補庫存處理，請告知時點。", color=GREY, after=8)

# 要求三
h1("要求三：「借料未還」可視狀態")
para("接受裁示：不用中性的「退料在途」，改用帶催促語氣的「借料未還」——一句給倉管/主管看了會"
     "警覺的警語（「還沒還，是不是該問一下？」）。判定 = TL 零件已由倉管 repair-pick 正式出庫、"
     "尚未全部歸還回庫；天數從實際出庫日起算（與 due_by 語意不同）。")
h2("門檻建議（依現有 UI 色階規範）")
para("Russell 指定「3 天內正常 / 超過 7 天警示」。建議中段補一階 amber「注意」做漸進升級（要收斂"
     "回兩階，拿掉 amber 即可）：")
table(["分級", "天數", "語氣", "UI 色票"], [
    ["資訊", "≤ 3 天", "純資訊提示", "資訊藍 #185FA5 / #EAF4FB"],
    ["注意", "4 – 7 天", "開始注意", "警告 amber #854F0B / #FDF3E3"],
    ["警示", "> 7 天", "主管關注（加 ⚠）", "危險紅 #CC0000 / #FDECEA"],
])
para("門檻已抽成可調整常數（INFO_MAX=3、WARNING_MIN=7），確認後一行可改。")
h2("UI 呈現與位置")
para("位置：TL 工單 detail 標題列 badge 區，緊接既有「今日 HH:MM 截止」badge，SA / 主管一開工單就看到。")
img("09_borrowed-not-returned-closeup.png", 6.5,
    "圖 4：真實工單 TL-IN-260617-001 的 badge 列 —— 「今日 18:00 截止」（紅）＋「借料未還 0 天」（藍，資訊階，因剛出庫）。amber / 紅階會在真實借料跨 4 天 / 7 天時自動顯示。")
para("為何先做 TL detail：這是「已出庫但還沒結案」風險窗最直接、改動最小、主管最常開的地方。"
     "倉管端的退料待確認（return-in）已有自己的到期/逾期紅黃分級，涵蓋「結案後等倉管點收」那段窗。"
     "若要把「借料未還」也鋪到 return-in 列 / 工單列表，可下一輪追加（共用同一套門檻邏輯）。", color=GREY)
para("我方刻意不用 SQL 竄改出庫日去「擺拍」紅階截圖（那正是上一輪被指出的塞資料問題）；分級邏輯"
     "由上表色票與門檻常數明確界定、可被審。", color=GREY, after=8)

h1("等裁示")
bullet("「借料未還」門檻採三階（含 amber）或收斂回兩階？")
bullet("是否把「借料未還」chip 也鋪到 return-in 列 / TL 工單列表（下一輪追加）？")
bullet("真實 e2e 建立的 TL-IN-260617-001（已實際出庫 2 件）清除時點。")

doc.save(OUT)
print("WROTE", OUT)
